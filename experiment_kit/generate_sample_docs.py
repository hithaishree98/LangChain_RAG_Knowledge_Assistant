"""
Generates Q3_QBR_Notes.docx and Meridian_SOW_v2.pdf from scratch.
Run from the repo root: python experiment_kit/generate_sample_docs.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

OUT = Path(__file__).parent / "sample_docs"
OUT.mkdir(exist_ok=True)


# ── DOCX ─────────────────────────────────────────────────────────────────────

def make_docx():
    doc = Document()

    def h1(text):
        doc.add_heading(text, level=1)

    def h2(text):
        doc.add_heading(text, level=2)

    def h3(text):
        doc.add_heading(text, level=3)

    def p(text):
        doc.add_paragraph(text)

    def table_rows(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = h
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = val

    def numbered(items):
        for i, item in enumerate(items, 1):
            doc.add_paragraph(f"{i}. {item}", style="List Number")

    h1("Q3 2024 Quarterly Business Review — Meridian Financial")
    p("Date: 2024-09-10  |  Format: Video call  |  Prepared by: Elena Martinez, VP Customer Success")

    h1("Attendees")
    table_rows(
        ["Name", "Role", "Company"],
        [
            ["Sarah Chen", "Chief Technology Officer", "Meridian Financial"],
            ["Mike Rodriguez", "Head of Data Platform", "Meridian Financial"],
            ["Priya Desai", "Analytics Lead", "Meridian Financial"],
            ["David Park", "Forward Deployed Engineer", "Orion Analytics"],
            ["Elena Martinez", "VP Customer Success", "Orion Analytics"],
        ],
    )

    h1("Account Health")

    h2("Seat Utilization")
    p(
        "Meridian currently has 41 of 50 contracted analyst seats active, distributed across four "
        "offices: 14 in New York City, 11 in Denver, 9 in Seattle, and 7 in Los Angeles. Seat "
        "utilization is at 82%, which is within the healthy range but approaching the ceiling of "
        "the contracted allocation."
    )
    p(
        "Priya Desai indicated that the team expects to request 20 to 30 additional seats by the "
        "end of Q1 2025 as the Denver and Seattle offices grow. At the contracted overage rate of "
        "$180 per seat per month, 25 additional seats would add $4,500 to the monthly bill, and "
        "30 additional seats would add $5,400. Sarah Chen confirmed Q1 2025 as the likely timing "
        "for a formal expansion request."
    )

    h2("Platform Adoption")
    p(
        "The regulatory reporting workflow is used daily by all 41 active analysts and is the "
        "primary driver of platform value for Meridian. Customer segmentation is used by 12 of "
        "41 seats, predominantly in the NYC office. Real-time transaction analytics is used by "
        "only 3 analysts. Priya Desai noted that low adoption of the analytics feature is due to "
        "a missing Snowflake export format — the team requested a Parquet-compatible schema in Q2 "
        "and the request is still open with the Orion product team."
    )

    h2("Event Volume")
    p(
        "Current ingest is approximately 4,200 events per second on average, peaking at 6,800 "
        "during market open between 09:30 and 10:00 ET. On a monthly basis Meridian is ingesting "
        "approximately 142 million events against a contracted limit of 200 million events per "
        "month, leaving 29% headroom before overage charges apply. Overage is billed at $0.08 "
        "per 10,000 events above the 200 million threshold. If the planned seat expansion "
        "proceeds and analyst activity increases proportionally, monthly event volume could "
        "approach 170-180 million events, still within the contracted limit."
    )

    h1("Open Issues and Action Items")

    h2("Login Latency (TICK-4521)")
    p(
        "Meridian is experiencing elevated login latency for analyst seats. P95 login time has "
        "degraded from approximately 800ms to approximately 2,400ms over a three-week window "
        "ending September 14. The degradation is worst in west coast offices (Denver, Seattle, "
        "LA) during the 09:00-10:00 ET window. Root cause identified: the load balancer "
        "sticky-session rule pins users to a backend node that saturates under Meridian's bursty "
        "morning traffic before the 15-minute rebalance window expires."
    )
    p(
        "David Park committed to delivering a production fix by September 25, 2024, with daily "
        "status updates to Sarah Chen until the fix ships. The fix reduces the sticky-session "
        "timeout from 15 minutes to 3 minutes and adds a CPU-pressure trigger to force-evict "
        "idle sessions when node CPU exceeds 80%. Staging validation is scheduled for "
        "September 20 with Mike Rodriguez's team observing."
    )

    h2("EU Region Availability")
    p(
        "Elena Martinez confirmed the GA target for the eu-west-2 region is April 2025. If "
        "Meridian's UK subsidiary launches in Q2 2025, they would be right at the edge of "
        "Orion's regional availability. Elena noted the risk of slippage is moderate, based on "
        "prior region launch history, and declined to make a firm commitment."
    )
    p(
        "David Park agreed to produce a UK contingency feasibility document by October 31, 2024. "
        "The document will cover the architecture for running UK operations on us-east-2 during "
        "any gap before eu-west-2 GA, including a data replication pipeline, latency tradeoffs "
        "for UK users, and estimated additional infrastructure cost. Sarah Chen stated she needs "
        "options, not just the primary plan."
    )

    h2("Action Items")
    numbered([
        "David Park: deliver login latency production fix by 2024-09-25 with daily status updates to Sarah Chen",
        "David Park: deliver UK contingency feasibility document by 2024-10-31",
        "Elena Martinez: deliver expanded seat count proposal with pricing for 20-30 additional seats by 2024-11-15",
        "Priya Desai: provide Snowflake export format specification to Orion product team by 2024-09-30",
    ])

    h1("Expansion Opportunities")

    h2("Seat Expansion")
    p(
        "Current seat utilization of 82% (41 of 50) is at the low end of the upsell window for "
        "enterprise SaaS accounts. Industry benchmarks suggest 80-90% utilization signals "
        "imminent expansion demand. Meridian's Denver and Seattle teams are growing and are the "
        "primary driver of the expected Q1 2025 expansion request."
    )
    p(
        "A 25-seat expansion at $180 per seat per month generates $4,500 per month in incremental "
        "revenue, or $54,000 annualized. A 30-seat expansion generates $5,400 per month, or "
        "$64,800 annualized. Sarah Chen indicated the ask is likely 20 to 30 additional seats, "
        "with the final number depending on hiring pace in the west coast offices."
    )

    h2("Reranker Upgrade")
    p(
        "Meridian's tenant has the cross-encoder reranker disabled, which is the default for the "
        "Enterprise tier. Enabling the reranker has shown a measured 12% improvement in "
        "recall@5 on Meridian's document corpus in September benchmarking. The reranker is "
        "included in the Enterprise tier at no additional charge and can be enabled with a single "
        "configuration change. David Park recommends enabling it before the UK launch to ensure "
        "retrieval quality across a larger multi-region document corpus."
    )

    h2("EU Region Launch")
    p(
        "When eu-west-2 reaches GA in April 2025, Meridian would be an early adopter. Adding one "
        "additional region costs $4,000 per month per the SOW Section 5.2 overage terms. For the "
        "UK subsidiary launch this would add $4,000 per month to the platform fee, bringing the "
        "total to $32,000 per month before any seat expansion. If the 25-seat expansion also "
        "proceeds, the total monthly fee would be $36,500."
    )

    h1("Risk Assessment")

    h2("Login Latency Risk")
    p(
        "The login latency issue (TICK-4521) represents the most acute risk to the account. "
        "Sarah Chen's CFO has directly questioned the $28,000 per month platform fee given the "
        "degradation. A slip past the September 25 production fix deadline will trigger executive "
        "escalation and materially increase renewal risk ahead of the 2026-06-30 contract expiry. "
        "The account health score has been downgraded from Green to Amber by the CS team pending "
        "resolution. If the fix ships by September 25 and the QBR action items are completed on "
        "schedule, the account is expected to return to Green in October."
    )

    h2("Salesforce Connector Risk")
    p(
        "The Salesforce connector v2.4.1 has a documented bulk update limit of 50,000 records "
        "per hour. Meridian's current volume is approximately 12,000 records per hour, well "
        "within the limit. However, if the planned seat expansion proceeds and analyst activity "
        "increases event output by 40-50%, volume could approach 17,000-18,000 records per hour. "
        "This is still within the v2.4.1 limit, but the trend should be monitored monthly via "
        "the sfdc_bulk_throughput dashboard."
    )
    p(
        "The v3.x Salesforce connector, which removes the bulk update limit entirely, is targeted "
        "for availability in Q1 2025. David Park recommends planning the connector migration "
        "alongside the seat expansion to avoid approaching the v2.4.1 ceiling. Migration is a "
        "configuration change and does not require a maintenance window."
    )

    h2("Plaid Schema Migration Risk")
    p(
        "Plaid has announced end-of-life for schema v3.2.1 in Q2 2025. Orion's v4.0 connector "
        "support is targeted for Q1 2025. If the v4.0 connector slips past Q1, Meridian will "
        "be operating on a deprecated schema with no upgrade path during a window that coincides "
        "with the planned UK subsidiary launch. This is rated as a medium risk. Mitigation: "
        "Meridian has requested early beta access to the v4.0 connector (TICK-4702) and should "
        "plan to begin the migration process in January 2025 regardless of GA status, using "
        "the compatibility shim as a fallback. The five-step migration process is documented in "
        "the Orion Integration Guide Section 4.1."
    )

    h2("EU Region Availability Risk")
    p(
        "The eu-west-2 GA target of April 2025 carries moderate slippage risk, as acknowledged "
        "by Elena Martinez on the September 15 call. Meridian's UK subsidiary is planning a Q2 "
        "2025 launch, leaving minimal buffer if the region slips. The contingency plan — running "
        "UK operations on us-east-2 with a data replication pipe — introduces latency for UK "
        "users and may have compliance implications under UK GDPR regarding data residency. "
        "David Park will produce a feasibility document covering these tradeoffs by October 31. "
        "The compliance question has been flagged to the Orion legal team for review."
    )

    h1("Historical Performance Review")

    h2("Q1 2024 Performance")
    p(
        "Q1 2024 was Meridian's first full quarter on the Orion platform following the Q3 2023 "
        "onboarding and Q4 2023 pilot. The platform processed 118 million events in Q1, below "
        "the 200 million monthly average limit and well within contracted capacity. Auth p95 "
        "latency averaged 790ms across the quarter, consistent with the SLA target. The regulatory "
        "reporting workflow was deployed in February 2024 following two weeks of UAT with Priya "
        "Desai's team. No P0 incidents were recorded in Q1. One P1 incident (TICK-4389) related "
        "to a Snowflake export delay was resolved within the 1-business-day SLA."
    )

    h2("Q2 2024 Performance")
    p(
        "Q2 2024 saw the highest event volume to date: 138 million events in April, 141 million "
        "in May, and 144 million in June, reflecting growth in Meridian's customer base and "
        "increased analyst adoption of the regulatory reporting workflow. Seat utilization grew "
        "from 74% (37 of 50 seats) at the start of Q2 to 80% (40 of 50 seats) by end of June. "
        "The customer segmentation workflow was enabled for the NYC analyst team in May. No P0 "
        "incidents occurred in Q2. The duplicate invoice incident (TICK-4602) was resolved in "
        "September 2024 but its root cause — a Stripe webhook replay — originated from a billing "
        "cycle processed in Q2; the hotfix was deployed retroactively. Auth p95 latency remained "
        "below 850ms throughout Q2."
    )

    h2("Q3 2024 Performance (through September 14)")
    p(
        "Q3 2024 introduced the first significant operational challenge: the login latency "
        "degradation documented in TICK-4521, which began September 3 and remained open at the "
        "time of the QBR on September 10. Event volume for July was 140 million and for August "
        "was 142 million, continuing the growth trend. The duplicate invoice incident (TICK-4602) "
        "was discovered and resolved on September 1. At the time of the QBR, 41 of 50 analyst "
        "seats are active (82% utilization). The Snowflake export workflow is operating within "
        "the 8-minute lag SLA. The Salesforce sync is running at 12,000 records per hour, well "
        "within the v2.4.1 50,000 records per hour limit."
    )

    h1("Financial Analysis")

    h2("Current Monthly Spend")
    table_rows(
        ["Line Item", "Monthly Amount", "Notes"],
        [
            ["Base platform fee (Enterprise)", "$28,000", "50 seats, 200M events"],
            ["Professional services (FDE hours)", "Included", "20 hrs/mo included per SOW 5.3"],
            ["Overage events", "$0", "142M of 200M limit used"],
            ["Additional regions", "$0", "Only us-east-2 + us-west-2 DR"],
            ["Total current monthly", "$28,000", ""],
        ],
    )

    h2("Projected Monthly Spend After Expansion")
    p(
        "If Meridian proceeds with the planned 25-seat expansion and adds the eu-west-2 region "
        "for the UK subsidiary launch, the new monthly fee structure would be as follows. The "
        "base platform fee remains $28,000. The 25 additional analyst seats at $180 per seat "
        "per month add $4,500 to the monthly total. The eu-west-2 region adds $4,000 per month "
        "per SOW Section 5.2. The combined total would be $36,500 per month, representing a "
        "30.4% increase over the current fee. Annualized, this is $438,000 compared to the "
        "current $336,000. The seat expansion alone (without EU region) would be $32,500 per "
        "month. The EU region alone (without seat expansion) would be $32,000 per month. Elena "
        "Martinez will provide a formal expansion proposal with final pricing by November 15."
    )

    h2("Return on Investment Analysis")
    p(
        "Meridian's finance team conducted an informal ROI analysis at the 12-month mark "
        "(August 2024). The regulatory reporting workflow has reduced the time for Meridian's "
        "compliance team to produce weekly regulatory reports from 3.5 analyst-days to 0.5 "
        "analyst-days, a saving of approximately 3 analyst-days per week across 41 analysts. "
        "At Meridian's loaded analyst cost of $220 per hour, this represents approximately "
        "$132,000 per month in labor cost avoidance — more than 4x the monthly platform fee. "
        "The customer segmentation workflow adopted by 12 NYC analysts has contributed to a "
        "measurable improvement in targeted campaign conversion rates, though the finance team "
        "has not yet formally quantified this contribution. The ROI analysis supports the "
        "planned seat expansion and reinforces Meridian's intent to renew at the 2026-06-30 "
        "expiry date, subject to resolution of the current operational issues."
    )

    h1("Q4 2024 and Q1 2025 Planning")

    h2("Priority Actions Q4 2024")
    numbered([
        "Close TICK-4521 (login latency fix) by September 25 — David Park leading",
        "Deliver UK contingency feasibility document by October 31 — David Park",
        "Deliver seat expansion proposal with pricing by November 15 — Elena Martinez",
        "Deliver Snowflake export format specification to Orion product team by September 30 — Priya Desai",
        "Schedule Plaid v4.0 migration planning session with Priya Desai and Orion product team — October or November",
        "Confirm EU region (eu-west-2) GA date by end of Q4 — Elena Martinez to track",
    ])

    h2("Priority Actions Q1 2025")
    numbered([
        "Execute 25-seat expansion (target January 2025) — Elena Martinez to contract, David Park to provision",
        "Begin Plaid v3.2.1 to v4.0 migration (target January 2025 start) — Priya Desai leading",
        "Upgrade Salesforce connector from v2.4.1 to v3.x (target February 2025) — David Park",
        "Enable cross-encoder reranker for Meridian tenant (target post-latency-fix) — David Park",
        "Prepare for eu-west-2 launch (target April 2025 GA) — David Park and Elena Martinez",
    ])

    path = OUT / "Q3_QBR_Notes.docx"
    doc.save(str(path))
    print(f"Saved: {path}")


# ── PDF ──────────────────────────────────────────────────────────────────────

class PDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 6, "Orion Analytics Platform -- Master Services Agreement  |  Confidential -- Meridian Financial Inc.", align="L")
        self.ln(4)

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 6, f"Document ID: ORN-MSA-2024-0871  |  Rev: 2.1  |  Page {self.page_no()} of {{nb}}  |  Effective: 2024-07-01", align="C")  # noqa: E501

    def section_title(self, text):
        self.ln(6)
        self.set_font("Helvetica", "B", 13)
        self.set_text_color(29, 54, 93)
        self.cell(0, 8, _a(text), ln=True)
        self.set_draw_color(29, 54, 93)
        self.line(self.get_x(), self.get_y(), self.get_x() + 170, self.get_y())
        self.ln(3)
        self.set_text_color(0, 0, 0)

    def sub_title(self, text):
        self.ln(4)
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(45, 74, 124)
        self.cell(0, 7, _a(text), ln=True)
        self.set_text_color(0, 0, 0)

    def body(self, text):
        self.set_font("Helvetica", "", 10)
        self.multi_cell(0, 5.5, _a(text))
        self.ln(2)

    def kv_table(self, rows):
        self.set_font("Helvetica", "", 10)
        col_w = [60, 110]
        for label, value in rows:
            x = self.get_x()
            y = self.get_y()
            self.set_fill_color(240, 244, 248)
            self.cell(col_w[0], 7, _a(label), border=1, fill=True)
            self.cell(col_w[1], 7, _a(value), border=1)
            self.ln()
        self.ln(3)

    def priority_table(self, headers, rows):
        self.set_font("Helvetica", "B", 9)
        self.set_fill_color(240, 244, 248)
        widths = [20, 80, 40, 30]
        for i, h in enumerate(headers):
            self.cell(widths[i], 7, _a(h), border=1, fill=True)
        self.ln()
        self.set_font("Helvetica", "", 9)
        for row in rows:
            for i, c in enumerate(row):
                self.cell(widths[i], 7, _a(c), border=1)
            self.ln()
        self.ln(3)

    def param_table(self, headers, rows):
        self.set_font("Helvetica", "B", 9)
        self.set_fill_color(240, 244, 248)
        widths = [55, 60, 55]
        for i, h in enumerate(headers):
            self.cell(widths[i], 7, _a(h), border=1, fill=True)
        self.ln()
        self.set_font("Helvetica", "", 9)
        for row in rows:
            for i, c in enumerate(row):
                self.cell(widths[i], 7, _a(c), border=1)
            self.ln()
        self.ln(3)


def _a(text: str) -> str:
    """Replace non-latin-1 chars so Helvetica doesn't choke."""
    return (text
            .replace("—", "--")   # em dash
            .replace("–", "-")    # en dash
            .replace("’", "'")    # right single quote
            .replace("‘", "'")    # left single quote
            .replace("“", '"')    # left double quote
            .replace("”", '"'))   # right double quote


def make_pdf():
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    # Cover
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(29, 54, 93)
    pdf.ln(8)
    pdf.cell(0, 10, "STATEMENT OF WORK", align="C", ln=True)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 8, "Orion Analytics Platform -- Enterprise Deployment", align="C", ln=True)
    pdf.ln(6)
    pdf.set_text_color(0, 0, 0)

    pdf.kv_table([
        ("Customer:", "Meridian Financial, Inc."),
        ("Customer ID:", "CUST-0287"),
        ("Contract Term:", "24 months (2024-07-01 through 2026-06-30)"),
        ("Renewal Date:", "2026-06-30"),
        ("Account Owner (Orion):", "David Park, Forward Deployed Engineer"),
        ("Account Exec (Meridian):", "Sarah Chen, CTO"),
        ("Total Contract Value:", "$480,000 annual / $960,000 total"),
        ("Billing Cadence:", "Monthly invoice, NET-30"),
    ])

    pdf.section_title("1. Executive Summary")
    pdf.body(
        "This Statement of Work (SOW) governs the enterprise deployment of the Orion Analytics "
        "Platform for Meridian Financial, Inc. Orion provides real-time transaction analytics, "
        "customer segmentation, and regulatory reporting workflows. This agreement supersedes the "
        "pilot agreement signed 2024-02-15."
    )
    pdf.body(
        "Meridian operates across 14 U.S. states with approximately 2.4 million active customers. "
        "The platform will ingest transaction events from three upstream systems (Plaid, Stripe, "
        "and Meridian's internal ledger) at a projected peak throughput of 8,000 events per second."
    )

    pdf.section_title("2. Service Level Agreements")
    pdf.body(
        "Orion commits to the following availability and response SLAs. Credits accrue against "
        "the next invoice per Section 8 if thresholds are missed."
    )

    pdf.sub_title("2.1 Availability")
    pdf.body(
        "Uptime target: 99.9% monthly measured against the control plane and API surface. "
        "Scheduled maintenance windows (Sundays 02:00-04:00 UTC) are excluded from measurement."
    )

    pdf.sub_title("2.2 Incident Response")
    pdf.priority_table(
        ["Priority", "Definition", "First Response", "Resolution"],
        [
            ["P0", "Complete outage, data loss, security breach", "30 minutes", "4 hours"],
            ["P1", "Major functionality broken, no workaround", "2 hours", "1 business day"],
            ["P2", "Functionality degraded, workaround exists", "8 business hours", "5 business days"],
            ["P3", "Minor issue, feature request, cosmetic", "2 business days", "Best effort"],
        ],
    )
    pdf.body(
        "All P0 incidents trigger an automated page to the on-call engineer assigned to Meridian's "
        "account. The on-call engineer must acknowledge the page within 30 minutes and begin "
        "active investigation. Every page notification includes the incident ID, the affected "
        "service, and the time the threshold was breached."
    )
    pdf.body(
        "The P0 escalation chain uses explicit time thresholds measured from incident declaration: "
        "T+0 -- the on-call engineer is paged immediately upon incident declaration and must "
        "acknowledge within 30 minutes; T+15 -- if unacknowledged within 15 minutes of the "
        "initial page, the team lead is automatically notified; T+30 -- if the incident remains "
        "open 30 minutes from the original incident start, VP Engineering is notified directly; "
        "T+60 -- at 60 minutes from incident start the executive sponsor (Sarah Chen, CTO) is "
        "briefed by VP Engineering regardless of acknowledgment status. This chain applies to "
        "all P0 incidents and to any P1 that breaches the 2-hour first-response SLA."
    )

    pdf.sub_title("2.3 Service Credits")
    pdf.body(
        "Monthly uptime below 99.9% but at or above 99.0%: 10% credit of that month's fees. "
        "Below 99.0% but at or above 95.0%: 25% credit. Below 95.0%: 50% credit. Credits are "
        "capped at 50% of any single month's fees and applied to the following invoice."
    )

    pdf.section_title("3. Deployment Configuration")

    pdf.sub_title("3.1 Hosting and Region")
    pdf.body(
        "Primary region: us-east-2 (Ohio). Secondary region for disaster recovery: us-west-2 "
        "(Oregon). All workloads run across three availability zones in the primary region. "
        "Data residency: all customer data remains within U.S. territory per Meridian's "
        "compliance posture (SOC 2 Type II, PCI DSS Level 1)."
    )

    pdf.sub_title("3.2 Retrieval Engine Configuration")
    pdf.body("The Orion knowledge retrieval subsystem is configured with the following defaults for Meridian's tenant:")
    pdf.param_table(
        ["Parameter", "Default Value", "Tuning Range"],
        [
            ["Retrieval top-k", "k = 2", "1 - 10"],
            ["Chunk size / overlap", "1000 / 200 characters", "500-2000 / 50-400"],
            ["Embedding model", "nomic-ai/nomic-embed-text-v1.5", "Managed by Orion"],
            ["Reranker", "Disabled by default", "Available on Enterprise"],
            ["Hybrid search (BM25)", "Enabled", "Boolean toggle"],
            ["Query timeout", "8 seconds", "2 - 30 seconds"],
        ],
    )

    pdf.sub_title("3.3 Backup and Disaster Recovery")
    pdf.body(
        "Automated snapshots are taken every 24 hours and replicated to us-east-2 across 3 "
        "availability zones. A weekly cross-region snapshot is retained for 90 days in us-west-2. "
        "RTO: 4 hours. RPO: 24 hours. Restore drills are conducted quarterly on behalf of the customer."
    )

    pdf.section_title("4. Data Sources and Integrations")

    pdf.sub_title("4.1 Upstream Systems")
    pdf.body("Meridian will integrate the following upstream data sources during the rollout:")
    pdf.body(
        "- Plaid -- transaction feed via webhooks, event schema v3.2.1\n"
        "- Stripe -- payment events via the API, using API version 2024-04-10\n"
        "- Internal ledger -- nightly batch export via SFTP, pipe-delimited format"
    )

    pdf.sub_title("4.2 Downstream Integrations")
    pdf.body("Reports and scored segments are delivered to:")
    pdf.body(
        "- Meridian's Snowflake warehouse (EXTERNAL_STAGE:ORION_EXPORT) -- hourly\n"
        "- Salesforce (Customer_360 object) -- via the Orion-Salesforce connector v2.4.1\n"
        "- Regulatory reporting pipeline -- daily 06:00 UTC\n"
        "- Slack (#orion-alerts channel) -- real-time P0/P1 notifications"
    )

    pdf.sub_title("4.3 Known Integration Constraints")
    pdf.body(
        "The Salesforce connector v2.4.1 has a documented limitation on bulk object updates "
        "above 50,000 records per hour. Meridian's expected volume is ~12,000 records per hour, "
        "within limits. If this threshold is approached, Meridian should engage their Orion FDE "
        "to plan migration to the v3.x connector (available Q1 2025)."
    )

    pdf.section_title("5. Pricing and Commercial Terms")

    pdf.sub_title("5.1 Platform Fee")
    pdf.body(
        "Base platform fee: $28,000 per month (Enterprise tier). Includes 5 named admin users, "
        "50 analyst seats, unlimited read-only access, and up to 200 million events ingested per month."
    )

    pdf.sub_title("5.2 Overage and Usage")
    pdf.body(
        "Events above 200M/month: $0.08 per 10,000 events. Additional analyst seats beyond 50: "
        "$180 per seat per month. Additional regions (beyond the contracted us-east-2 + "
        "us-west-2): $4,000 per region per month."
    )

    pdf.sub_title("5.3 Professional Services")
    pdf.body(
        "Onboarding and initial integration: $60,000 fixed-price, completed Q3 2024. Ongoing "
        "FDE support: 20 hours per month included, $280 per hour thereafter. Custom connector "
        "development: quoted per SOW amendment."
    )

    pdf.sub_title("5.4 Renewal")
    pdf.body(
        "Agreement renews automatically for successive 12-month terms unless either party "
        "provides 90 days written notice prior to the renewal date. Notice must be delivered "
        "in writing to the account executive named on this SOW."
    )
    pdf.body(
        "Pricing at renewal will increase by the lesser of CPI (U.S. Consumer Price Index, "
        "12-month trailing) or 5%, assessed at the renewal anniversary and applied to the base "
        "platform fee only. Overage rates, professional services rates, and per-seat pricing "
        "are not subject to the automatic increase and will be renegotiated at renewal if "
        "either party requests a change."
    )

    pdf.section_title("6. Security, Privacy, and Compliance")

    pdf.sub_title("6.1 Certifications")
    pdf.body(
        "Orion maintains SOC 2 Type II (annual audit, most recent: 2024-05-22), "
        "ISO 27001:2022, PCI DSS Level 1 (as a service provider), and HIPAA compliance posture. "
        "Attestation reports are available to Meridian under NDA."
    )

    pdf.sub_title("6.2 Encryption")
    pdf.body(
        "Data at rest is encrypted using AES-256 with AWS KMS customer-managed keys (CMKs) "
        "scoped to Meridian's tenant. CMKs are rotated automatically on an annual schedule; "
        "manual rotation is available via the admin console in five steps and takes effect "
        "within 15 minutes of confirmation. Field-level encryption is applied to all fields "
        "tagged as PII in Meridian's data classification schema, using envelope encryption "
        "so that individual field keys can be revoked without re-encrypting the full dataset. "
        "Data in transit is protected by TLS 1.3 with enforced Perfect Forward Secrecy (PFS) "
        "cipher suites; TLS 1.2 and below are rejected at the API gateway."
    )

    pdf.sub_title("6.3 Access Control")
    pdf.body(
        "SSO via SAML 2.0 integrated with Meridian's Okta tenant. MFA enforced for all admin "
        "actions. Session timeout: 30 minutes idle. All admin actions are logged to the audit "
        "trail and retained for 7 years."
    )

    pdf.sub_title("6.4 Incident Disclosure")
    pdf.body(
        "Security incidents affecting Meridian data are disclosed within 24 hours of "
        "confirmation, in accordance with Section 7 of the Master Services Agreement and "
        "applicable regulatory requirements (GLBA, state breach notification laws)."
    )

    pdf.sub_title("6.5 Audit Log Policy")
    pdf.body(
        "The Orion platform maintains an immutable audit log of all actions that modify system "
        "state or access customer data. Actions logged include: all admin console actions, API "
        "key creation and revocation, data export requests, configuration changes to retrieval "
        "parameters, user provisioning and deprovisioning, and all authentication events "
        "including SSO assertions and MFA challenges."
    )
    pdf.body(
        "Access to the audit log is restricted to users with the admin-role permission assigned "
        "in Meridian's Okta tenant. Audit logs are retained for 7 years in immutable storage "
        "and cannot be deleted or modified by any user, including Orion staff, without a "
        "formal legal hold process. Export is available as CSV via the admin console under "
        "Audit > Export, or programmatically as JSON via the /v2/audit API endpoint with an "
        "admin-scoped API key."
    )

    pdf.section_title("7. Incident Management and Post-Incident Review")

    pdf.sub_title("7.1 Incident Classification")
    pdf.body(
        "All incidents are classified at creation time based on customer impact. The classification "
        "determines the response SLA, escalation chain, and post-incident review requirements. "
        "P0 incidents (complete outage, data loss, or security breach affecting customer data) "
        "require a post-incident review document delivered to the customer within 5 business days "
        "of resolution. The document must cover: incident timeline, root cause analysis, immediate "
        "remediation steps taken, and preventive measures implemented to prevent recurrence. "
        "P1 incidents (major functionality impaired, no adequate workaround) require a summary "
        "email to the customer primary contact within 2 business days of resolution. P2 and P3 "
        "incidents are documented in the ticket system but do not require customer-facing "
        "post-incident communications unless the customer specifically requests one."
    )

    pdf.sub_title("7.2 Post-Incident Review Requirements")
    pdf.body(
        "Post-incident reviews for P0 incidents must include five required sections: (1) timeline "
        "with timestamps accurate to the minute, covering detection, acknowledgment, first "
        "customer notification, mitigation, and resolution; (2) root cause analysis using a "
        "5-Why methodology to trace the proximate cause to its systemic origin; (3) customer "
        "impact statement quantifying the number of affected users, the duration of impact, and "
        "any financial or compliance consequences; (4) immediate remediation steps taken during "
        "the incident window to restore service; (5) preventive measures implemented after "
        "resolution, with the name of the owner, the implementation date, and a verification "
        "step confirming the measure is in place. Reviews are delivered via email to the customer "
        "CTO and primary technical contact. For Meridian, this means Sarah Chen and Mike Rodriguez."
    )

    pdf.sub_title("7.3 Service Credit Calculation")
    pdf.body(
        "Service credits for SLA breaches are calculated as follows. Monthly uptime is measured "
        "as the percentage of minutes in the calendar month during which the Orion control plane "
        "and API surface return successful responses for at least 95% of requests. Scheduled "
        "maintenance windows (Sundays 02:00-04:00 UTC) are excluded from the calculation. "
        "If monthly uptime falls below 99.9% but remains at or above 99.0%, a credit of 10% "
        "of that month's base platform fee is applied to the following invoice. Below 99.0% but "
        "at or above 95.0% earns a 25% credit. Below 95.0% earns a 50% credit. Credits are "
        "capped at 50% of any single month's fees and do not carry over. Credits are applied "
        "automatically — Meridian does not need to submit a credit request. The monthly uptime "
        "calculation is available in the admin console under Account > SLA Report."
    )

    pdf.sub_title("7.4 Escalation Contacts")
    pdf.kv_table([
        ("Orion FDE (primary):", "David Park — david.park@orion.io"),
        ("Orion CS (account):", "Elena Martinez — elena.martinez@orion.io"),
        ("Orion On-Call (P0):", "pagerduty.orion.io/meridian (24x7)"),
        ("Orion VP Engineering:", "On-call escalation at T+30 per SLA 2.2"),
        ("Meridian CTO:", "Sarah Chen — sarah.chen@meridian.com"),
        ("Meridian Tech Lead:", "Mike Rodriguez — mike.rodriguez@meridian.com"),
    ])

    pdf.section_title("8. Data Governance and Privacy")

    pdf.sub_title("8.1 Data Classification")
    pdf.body(
        "Meridian's data processed by Orion is classified into three tiers. Tier 1 (PII and "
        "Financial): includes customer names, account numbers, transaction amounts, and any field "
        "tagged as PII in Meridian's data classification schema. Tier 1 data is subject to "
        "field-level encryption, strict access controls, and audit logging for every read event. "
        "Tier 2 (Aggregated Analytics): includes scored segments, behavioral clusters, and "
        "aggregated transaction metrics where individual records cannot be reconstructed. Tier 2 "
        "data is encrypted at rest but does not require field-level encryption. Tier 3 (Operational "
        "Metadata): includes system logs, performance metrics, and configuration state. Tier 3 "
        "data is retained for 90 days and is not subject to the 7-year retention requirement "
        "that applies to audit logs."
    )

    pdf.sub_title("8.2 Data Residency")
    pdf.body(
        "All Meridian customer data is stored exclusively within United States territory. "
        "Primary storage is in AWS us-east-2 (Ohio). Disaster recovery storage is in AWS "
        "us-west-2 (Oregon). Cross-region replication for DR purposes is encrypted in transit "
        "using TLS 1.3. No Meridian customer data is transferred to or processed in any region "
        "outside the United States without written consent from Meridian's CTO. This posture "
        "satisfies Meridian's current compliance requirements under Gramm-Leach-Bliley Act (GLBA), "
        "PCI DSS Level 1, and SOC 2 Type II. When the eu-west-2 region is provisioned for the UK "
        "subsidiary, a separate data residency agreement will be required to govern data flows "
        "between US and EU regions, specifically to address UK GDPR and EU GDPR requirements. "
        "Orion's legal team will engage Meridian's legal team at least 60 days before the "
        "eu-west-2 provisioning date to finalize this agreement."
    )

    pdf.sub_title("8.3 Retention and Deletion")
    pdf.body(
        "Customer data retained by Orion is subject to the following retention schedule. "
        "Transaction event data (raw ingested events): retained for 24 months from ingestion "
        "date, then deleted from primary storage. Aggregate analytics output (scored segments, "
        "cluster assignments): retained for the duration of the contract plus 90 days, then "
        "deleted unless Meridian requests an extended hold. Audit log entries: retained for "
        "7 years in immutable storage per SOW Section 6.5. Backup snapshots: daily snapshots "
        "retained for 90 days; weekly cross-region snapshots retained for 12 months. Upon "
        "contract termination, Orion will provide Meridian with a data export package containing "
        "all Tier 1 and Tier 2 data in Parquet format within 30 days of the termination date. "
        "All Meridian data will be securely deleted from Orion's systems within 60 days of "
        "the export delivery, with a deletion certificate provided to Meridian upon request."
    )

    pdf.sub_title("8.4 Right to Audit")
    pdf.body(
        "Meridian has the right to audit Orion's compliance with this SOW and its data handling "
        "obligations once per calendar year, with 30 days written notice. Audits may be conducted "
        "by Meridian's internal audit team or by a third-party auditor agreed upon by both parties. "
        "Orion will provide the auditor with access to relevant system logs, configuration records, "
        "and security documentation. Orion's SOC 2 Type II report (most recent: 2024-05-22) is "
        "available to Meridian under NDA and may satisfy audit requirements without requiring a "
        "separate on-site audit. If the audit reveals a material breach of Orion's obligations, "
        "Meridian may engage the dispute resolution process described in Section 9 of the "
        "Master Services Agreement."
    )

    pdf.section_title("9. Appendix A -- Pricing Schedule")

    pdf.sub_title("A.1 Full Pricing Table")
    pdf.param_table(
        ["Item", "Unit", "Price"],
        [
            ["Base platform fee (Enterprise)", "Per month", "$28,000"],
            ["Analyst seat (over 50)", "Per seat per month", "$180"],
            ["Additional region", "Per region per month", "$4,000"],
            ["Overage events (over 200M/mo)", "Per 10,000 events", "$0.08"],
            ["FDE hours (over 20/mo included)", "Per hour", "$280"],
            ["Custom connector development", "Per SOW amendment", "Quoted"],
            ["Onboarding (one-time, complete)", "Fixed price", "$60,000"],
        ],
    )

    pdf.sub_title("A.2 Example Expansion Scenarios")
    pdf.body(
        "Scenario 1 -- Seat expansion only (25 additional seats): base fee $28,000 + 25 seats "
        "at $180/seat = $4,500 in additional seat fees = $32,500 per month total. Annualized: "
        "$390,000. Scenario 2 -- EU region only (no additional seats): base fee $28,000 + "
        "eu-west-2 region at $4,000/month = $32,000 per month total. Annualized: $384,000. "
        "Scenario 3 -- Full expansion (25 additional seats plus EU region): base fee $28,000 + "
        "seat fees $4,500 + region fee $4,000 = $36,500 per month total. Annualized: $438,000. "
        "Scenario 4 -- Maximum expansion (30 additional seats plus EU region): $28,000 + $5,400 "
        "+ $4,000 = $37,400 per month total. All scenarios assume no event overage and no "
        "additional professional services hours beyond the 20 included per month."
    )

    pdf.sub_title("A.3 Renewal Pricing Mechanism")
    pdf.body(
        "At each 12-month renewal, the base platform fee is adjusted by the lesser of the "
        "U.S. Consumer Price Index (12-month trailing, all urban consumers, not seasonally "
        "adjusted, as published by the U.S. Bureau of Labor Statistics at the renewal "
        "anniversary) or 5%, whichever is smaller. The adjustment applies to the base platform "
        "fee only -- it does not apply to per-seat overage pricing ($180/seat), per-region "
        "pricing ($4,000/region), event overage rates ($0.08 per 10,000), or professional "
        "services rates ($280/hour). Those rates are fixed for the initial contract term and "
        "will be renegotiated separately at renewal if either party requests a change. The "
        "renewal adjustment is calculated and invoiced automatically; Meridian does not need "
        "to take any action. The renewal invoice will include a line item showing the CPI "
        "index value used and the resulting adjustment amount."
    )

    pdf.section_title("10. Signatures")
    pdf.ln(4)
    pdf.body("Accepted and agreed by authorized signatories:")
    pdf.ln(8)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(85, 6, "For Orion Analytics, Inc.", ln=False)
    pdf.cell(85, 6, "For Meridian Financial, Inc.", ln=True)
    pdf.ln(12)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(85, 6, "Name: Elena Martinez", ln=False)
    pdf.cell(85, 6, "Name: Sarah Chen", ln=True)
    pdf.cell(85, 6, "Title: VP, Customer Success", ln=False)
    pdf.cell(85, 6, "Title: Chief Technology Officer", ln=True)
    pdf.cell(85, 6, "Date: 2024-06-28", ln=False)
    pdf.cell(85, 6, "Date: 2024-06-30", ln=True)

    path = OUT / "Meridian_SOW_v2.pdf"
    pdf.output(str(path))
    print(f"Saved: {path}")


if __name__ == "__main__":
    make_docx()
    make_pdf()
    print("Done.")
