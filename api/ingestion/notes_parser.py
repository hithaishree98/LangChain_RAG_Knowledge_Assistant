"""
ingestion/notes_parser.py — Parse meeting notes (.txt or .docx).

Meeting notes are free-form prose — FDE's own writeup after a call.
Unlike transcripts, they have no speaker labels. We treat them as
plain text and sentence-chunk them.
"""
import os
import logging
from typing import List

from langchain_core.documents import Document

_log = logging.getLogger(__name__)


def parse(file_path: str) -> List[Document]:
    """
    Parse a meeting notes file into Document chunks.
    Supports .txt and .docx formats.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return _parse_docx(file_path)
    else:
        return _parse_txt(file_path)


def _parse_txt(file_path: str) -> List[Document]:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    if not text.strip():
        return []
    # Split into paragraphs (double newline = paragraph break)
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    docs = []
    for i, para in enumerate(paragraphs):
        if len(para.split()) < 3:
            continue  # skip very short fragments
        docs.append(Document(
            page_content=para,
            metadata={
                "source": file_path,
                "doc_type": "notes",
                "chunk_index": i,
            }
        ))
    _log.info("notes_parsed file=%s paragraphs=%d",
              os.path.basename(file_path), len(docs))
    return docs


def _parse_docx(file_path: str) -> List[Document]:
    try:
        import docx
        d = docx.Document(file_path)
        paragraphs = [p.text.strip() for p in d.paragraphs if p.text.strip()]
        docs = []
        for i, para in enumerate(paragraphs):
            if len(para.split()) < 3:
                continue
            docs.append(Document(
                page_content=para,
                metadata={
                    "source": file_path,
                    "doc_type": "notes",
                    "chunk_index": i,
                }
            ))
        return docs
    except Exception as e:
        _log.warning("notes_docx_parse_failed falling back to text: %s", e)
        # Fallback: try extracting text with docx2txt
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            return _parse_txt_content(text, file_path)
        except Exception as e2:
            _log.error("notes_parse_failed: %s", e2)
            return []


def _parse_txt_content(text: str, source_path: str) -> List[Document]:
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    return [
        Document(
            page_content=para,
            metadata={"source": source_path, "doc_type": "notes", "chunk_index": i}
        )
        for i, para in enumerate(paragraphs)
        if len(para.split()) >= 3
    ]
