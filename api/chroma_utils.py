import os
import re
import json
import logging
import functools
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Any, Callable, Dict, List, Optional, Tuple

_log = logging.getLogger(__name__)

from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_core.retrievers import BaseRetriever

# ── EXPERIMENT KNOBS ─────────────────────────────────────────────────────────
CHUNKING_MODE = os.getenv("CHUNKING_MODE", "full").lower()
RETRIEVAL_MODE = os.getenv("RETRIEVAL_MODE", "full").lower()
if CHUNKING_MODE != "full" or RETRIEVAL_MODE != "full":
    _log.info("experiment_mode_enabled chunking=%s retrieval=%s",
              CHUNKING_MODE, RETRIEVAL_MODE)

os.environ.setdefault("ANONYMIZED_TELEMETRY", "False")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CHROMA_DIR = os.getenv("CHROMA_DB_DIR") or os.path.join(BASE_DIR, "data", "chroma_db")
os.makedirs(CHROMA_DIR, exist_ok=True)

_embedder = None
_embedding_provider = None  # set by get_embedder() — "openai" | "huggingface"


def get_embedder() -> Embeddings:
    """Return the configured embedding model. Cached per-process.

    Two-tier selection:
      1. If OPENAI_API_KEY is set, use OpenAI's text-embedding-3-small
         (1536-dim, ~50-200ms per chunk over the network, ~$0.02 / 1M tokens).
         Right answer for production: no local model, no GPU needed, fast on
         any CPU container.
      2. Otherwise fall back to sentence-transformers/all-MiniLM-L6-v2
         (384-dim, ~150-400ms per chunk on CPU, free, no network). Right
         answer for offline dev / pytest / air-gapped deployments.

    The previous default (nomic-embed-text-v1.5, 440MB) has been removed —
    on CPU it dominated upload latency and the OpenAI/MiniLM split covers
    every realistic deployment without that cost.
    """
    global _embedder, _embedding_provider
    if _embedder is not None:
        return _embedder

    if os.getenv("OPENAI_API_KEY"):
        from langchain_openai import OpenAIEmbeddings
        _embedder = OpenAIEmbeddings(
            model="text-embedding-3-small",
            api_key=os.getenv("OPENAI_API_KEY"),
        )
        _embedding_provider = "openai"
    else:
        try:
            import torch
            device = "cuda" if torch.cuda.is_available() else "cpu"
        except ImportError:
            device = "cpu"
        _embedder = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": device},
            encode_kwargs={"normalize_embeddings": True},
        )
        _embedding_provider = "huggingface"

    _log.info("embedder_initialized provider=%s", _embedding_provider)
    return _embedder


def get_embedding_provider() -> str:
    """Return the active embedding provider name. Useful for /health and debug."""
    if _embedding_provider is None:
        # Trigger lazy initialization so callers get an honest answer
        get_embedder()
    return _embedding_provider or "unknown"


# The cache key intentionally includes a "provider tag" string so that, if
# the embedder ever swaps mid-process (e.g. an OPENAI_API_KEY toggle, or a
# warmup-time fallback to MiniLM), cached vectors of the wrong dimension
# never get returned for a query. Two embedders' vectors live in disjoint
# cache namespaces.
@functools.lru_cache(maxsize=512)
def _embed_cached_inner(provider_tag: str, text: str) -> tuple:
    return tuple(get_embedder().embed_query(text))


def embed_cached(text: str) -> tuple:
    """LRU-cached embedding for the active provider.

    Public surface kept identical to before (single ``text`` argument). The
    provider tag is read from :func:`get_embedding_provider` at call time so
    swapping providers automatically isolates the cache.
    """
    # Force lazy embedder init so get_embedding_provider returns the real value
    # rather than "unknown" on the first call.
    get_embedder()
    return _embed_cached_inner(_embedding_provider or "unknown", text)


class _LazyEmbedder(Embeddings):
    def embed_documents(self, texts): return get_embedder().embed_documents(texts)
    def embed_query(self, text):      return get_embedder().embed_query(text)


vectorstore = Chroma(collection_name="child_chunks", persist_directory=CHROMA_DIR,
                     embedding_function=_LazyEmbedder(),
                     collection_metadata={"hnsw:space": "cosine"})
parent_store = Chroma(collection_name="parent_chunks", persist_directory=CHROMA_DIR,
                      embedding_function=_LazyEmbedder(),
                      collection_metadata={"hnsw:space": "cosine"})


# ── Cross-encoder reranker ────────────────────────────────────────────────────
_cross_encoder = None


def _get_cross_encoder():
    global _cross_encoder
    if _cross_encoder is None:
        from sentence_transformers.cross_encoder import CrossEncoder
        _cross_encoder = CrossEncoder("cross-encoder/ms-marco-MiniLM-L-6-v2")
    return _cross_encoder


def rerank_docs(query: str, docs: List[Document], top_k: int = 4) -> List[Document]:
    if len(docs) <= top_k:
        return docs
    ce = _get_cross_encoder()
    pairs = [(query, doc.page_content) for doc in docs]
    scores = ce.predict(pairs)
    ranked = sorted(zip(scores, docs), key=lambda x: x[0], reverse=True)
    return [doc for _, doc in ranked[:top_k]]


def _diversity_cap(docs: List[Document], max_per_source: int = 2) -> List[Document]:
    """Cap chunks per source file so no single document dominates the answer.

    Preserves rerank ordering — keeps the first max_per_source chunks per
    filename and discards extras. Applied after reranking so priority is
    maintained within each source.
    """
    seen: Dict[str, int] = {}
    result = []
    for doc in docs:
        src = doc.metadata.get("filename") or doc.metadata.get("source") or ""
        count = seen.get(src, 0)
        if count < max_per_source:
            seen[src] = count + 1
            result.append(doc)
    return result


def warmup_models() -> None:
    """Force-load the embedding model, cross-encoder, and spaCy at API startup.

    Without this, models cold-load on the first request — spaCy in particular
    adds 2-3 seconds to the first PDF/DOCX upload. Pre-warming here ensures
    /health only reports ready once all models are actually usable.
    """
    # Trigger embedder load and run one embed call so lazy transformer
    # weights are fully initialized (not just the constructor).
    get_embedder().embed_query("warmup")
    _get_cross_encoder().predict([("warmup", "warmup")])
    # Pre-load spaCy so the first upload doesn't block while the model downloads.
    try:
        import spacy
        spacy.load("en_core_web_sm")
        _log.info("warmup_spacy_loaded")
    except Exception as e:
        _log.warning("warmup_spacy_failed error=%s (will lazy-load on first request)", str(e))


# ── BM25 ──────────────────────────────────────────────────────────────────────
def _tokenize(text): return re.findall(r"[a-z0-9]+", text.lower())

# Cache: user_id → (doc_count, BM25Okapi_index, texts, metas)
# Invalidated automatically when doc_count changes (upload / delete).
_bm25_cache: Dict[str, Tuple[int, Any, List, List]] = {}


def bm25_search(query, user_id, k=10):
    try:
        from rank_bm25 import BM25Okapi
    except ImportError:
        # Soft-degrade if rank_bm25 isn't installed — dense retrieval still works.
        # Logged once so the absence of BM25 is visible in deployment health checks.
        _log.warning("bm25_unavailable: rank_bm25 not installed; dense-only retrieval")
        return []
    try:
        result = vectorstore._collection.get(
            where={"user_id": {"$eq": user_id}},
            include=["documents", "metadatas"])
    except Exception as e:
        # A Chroma failure here means the vector store is unhealthy — caller
        # should NOT treat the empty list as "no BM25 hits." Logging the
        # exception so a health check / log scraper can detect the failure.
        # Returning [] keeps RRF working if dense still has results, but the
        # llm_breaker / health endpoint will surface the underlying issue.
        _log.error("bm25_search_chroma_failed", extra={"user_id": user_id, "error": str(e)})
        return []
    texts = result.get("documents") or []
    metas = result.get("metadatas") or []
    if not texts:
        return []

    # Reuse a cached index when the corpus size hasn't changed.
    # This avoids rebuilding an O(n) BM25 index from scratch on every query.
    cached = _bm25_cache.get(user_id)
    if cached is not None and cached[0] == len(texts):
        idx, texts, metas = cached[1], cached[2], cached[3]
    else:
        idx = BM25Okapi([_tokenize(t) for t in texts])
        _bm25_cache[user_id] = (len(texts), idx, texts, metas)

    scores = idx.get_scores(_tokenize(query))
    top_n = min(k, len(texts))
    top_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:top_n]
    return [Document(page_content=texts[i], metadata=metas[i] or {}) for i in top_indices]


def _rrf_doc_key(doc):
    """Stable identity key for a Document, used to merge dense + sparse rankings.

    Previous implementation used ``doc.page_content[:120]`` which collided
    when chunks shared a 120-char prefix. Real-world cases that hit this:
      - Contextual-retrieval prefix (`[Context: ...]`) made many chunks share
        their first 120 chars.
      - Transcript chunks of the same group with overlap shared opener turns.
      - Ticket sections often start with the same templated header.

    Prefer metadata identity when available (file_id + parent_chunk_id, or
    chunk_id, or filename + a content hash of full content). Fall back to a
    full-content hash as a last resort — never just the prefix.
    """
    md = getattr(doc, "metadata", {}) or {}
    file_id = md.get("file_id")
    parent_id = md.get("parent_chunk_id")
    chunk_id = md.get("chunk_id")
    if file_id is not None and parent_id:
        return ("p", file_id, parent_id)
    if file_id is not None and chunk_id:
        return ("c", file_id, chunk_id)
    # Worst case: hash the full content. Imports kept inline so the function
    # has no module-level deps beyond hashlib (already used elsewhere here).
    import hashlib
    h = hashlib.md5(
        (doc.page_content or "").encode("utf-8"), usedforsecurity=False
    ).hexdigest()
    return ("h", md.get("filename", ""), h)


def rrf_merge(dense, sparse, k=60):
    """Reciprocal Rank Fusion merge of two ranked Document lists.

    A document appearing in both rankings sees its scores summed via the
    standard RRF formula ``1 / (k + rank + 1)``. Identity is determined by
    :func:`_rrf_doc_key` — see that function for the collision-avoidance
    rationale.
    """
    scores = {}
    doc_map = {}
    for rank, doc in enumerate(dense):
        key = _rrf_doc_key(doc)
        scores[key] = scores.get(key, 0.0) + 1.0 / (k + rank + 1)
        doc_map[key] = doc
    for rank, doc in enumerate(sparse):
        key = _rrf_doc_key(doc)
        scores[key] = scores.get(key, 0.0) + 1.0 / (k + rank + 1)
        # Don't overwrite — keep the first-seen Document for stability.
        doc_map.setdefault(key, doc)
    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    return [doc_map[key] for key, _ in ranked]


def fetch_parents(child_chunks, user_id=None):
    parent_ids = []
    for doc in child_chunks:
        pid = doc.metadata.get("parent_chunk_id")
        if pid and pid not in parent_ids:
            parent_ids.append(pid)
    if not parent_ids:
        # Children have no parent_chunk_id metadata — this is the flat-mode
        # case (baseline / sentence). Caller treats children as the LLM context.
        return child_chunks
    if not user_id:
        # Hard guard against cross-tenant data exposure: without a user_id we
        # CANNOT safely query parent_store. Falling back to children is safe
        # (they're already user-scoped from the dense search). Surfaced so
        # callers don't mistake degraded context for normal operation.
        _log.warning("fetch_parents_skipped: no user_id provided; returning child chunks")
        return child_chunks
    try:
        result = parent_store._collection.get(
            ids=parent_ids,
            where={"user_id": {"$eq": user_id}},
            include=["documents", "metadatas"],
        )
        docs = result.get("documents") or []
        metas = result.get("metadatas") or []
        if docs:
            return [Document(page_content=docs[i], metadata=metas[i] or {}) for i in range(len(docs))]
        # Parent IDs didn't resolve — workspace may have been wiped between
        # child upload and query. Falling back to children is the only
        # recoverable option, but logged so operators can investigate.
        _log.warning("fetch_parents_empty_result", extra={
            "user_id": user_id, "parent_id_count": len(parent_ids)})
    except Exception as e:
        # parent_store unhealthy — log loudly so health checks / log scrapers
        # can detect the degraded state. The brief still completes (with
        # smaller child-chunk context) but quality will be lower until fixed.
        _log.error("fetch_parents_failed", extra={
            "user_id": user_id, "error": str(e)})
    return child_chunks


# ── Hybrid retriever with RETRIEVAL_MODE switch ───────────────────────────────
class HybridRetriever(BaseRetriever):
    user_id: str = "default"
    k_dense: int = 10
    k_bm25: int = 10
    k_rerank: int = 6

    class Config:
        arbitrary_types_allowed = True

    def _get_relevant_documents(self, query: str) -> List[Document]:
        _dense_retriever = vectorstore.as_retriever(
            search_kwargs={"k": self.k_dense, "filter": {
                "$and": [
                    {"user_id": {"$eq": self.user_id}},
                    {"is_latest_version": {"$eq": 1}},
                ]
            }}
        )
        try:
            dense = _dense_retriever.invoke(query)
        except AttributeError:
            dense = _dense_retriever.get_relevant_documents(query)

        if RETRIEVAL_MODE == "dense":
            # Dense only, top-k_rerank results
            return self._maybe_boost(query, _diversity_cap(dense[:self.k_rerank]))

        sparse = bm25_search(query, self.user_id, k=self.k_bm25)
        merged = rrf_merge(dense, sparse)

        if RETRIEVAL_MODE == "dense_bm25":
            # Dense + BM25 via RRF, no reranker, no parent fetch
            return self._maybe_boost(query, _diversity_cap(merged[:self.k_rerank]))

        # Full pipeline
        child_top = rerank_docs(query, merged, top_k=self.k_rerank)
        child_top = _diversity_cap(child_top)
        parents = fetch_parents(child_top, self.user_id)
        return self._maybe_boost(query, parents)

    def _maybe_boost(self, query: str, docs: List[Document]) -> List[Document]:
        """Re-rank by recency only when the query has a time-sensitive intent.

        Without the keyword guard, every retrieval would skew toward newer docs
        even for purely topical queries, hurting recall. With the guard, queries
        like "what's the most recent agreement" correctly surface newer chunks
        while "what's the SLA" stays semantically ranked.
        """
        if not _wants_recency(query) or not docs:
            return docs
        return _recency_boost(docs)

    async def _aget_relevant_documents(self, query):
        return self._get_relevant_documents(query)


def get_retriever_for_user(user_id):
    return HybridRetriever(user_id=user_id)


# ── PDF / DOCX / HTML loaders (with sentence-aware path) ─────────────────────
def noise_filter(text):
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped: continue
        if re.fullmatch(r"[-–—]\s*\d+\s*[-–—]?|\d+\s*[-–—]", stripped): continue  # require at least one dash so bare years ("2024") are preserved
        if re.fullmatch(r"[Pp]age\s+\d+(\s+of\s+\d+)?", stripped): continue
        if len(stripped.split()) < 4 and len(stripped) < 40:
            # Preserve section headers ("Risks:", "Summary:") and bullet points
            is_header = stripped.endswith(":")
            is_bullet = stripped[:1] in ("•", "-", "*", "–", "○", "▸", "·")
            if not is_header and not is_bullet:
                continue
        cleaned.append(stripped)
    return "\n".join(cleaned)


def sentence_chunk(text, size=256, overlap=64):
    sentences = _split_sentences(text)
    if not sentences: return []
    chunks = []
    i = 0
    while i < len(sentences):
        group = []; word_count = 0; j = i
        while j < len(sentences) and word_count < size:
            group.append(sentences[j])
            word_count += len(sentences[j].split())
            j += 1
        chunk_text = " ".join(group).strip()
        if chunk_text:
            chunks.append(Document(page_content=chunk_text))
        # Break immediately when all sentences are consumed — continuing would
        # produce tiny tail artifact chunks (same bug as transcript_chunker).
        if j >= len(sentences):
            break
        overlap_words = 0; step = j
        while step > i + 1 and overlap_words < overlap:
            step -= 1
            overlap_words += len(sentences[step].split())
        i = max(i + 1, step)
    return chunks


# Module-level spaCy cache. Loading the model is ~500ms; without caching,
# every PDF / plain-text chunk parse paid that cost. With this, the first
# call loads it and every subsequent _split_sentences uses the same nlp
# pipeline. Set to False if spaCy is unavailable so we don't keep retrying.
_nlp = None
_spacy_unavailable = False


def _split_sentences(text):
    global _nlp, _spacy_unavailable
    if _spacy_unavailable:
        # Regex fallback — already proven by the original code path.
        parts = re.split(r"(?<=[.!?])\s+", text)
        return [p.strip() for p in parts if p.strip()]
    if _nlp is None:
        try:
            import spacy
            _nlp = spacy.load("en_core_web_sm", disable=["ner", "tagger", "parser"])
            try:
                _nlp.enable_pipe("senter")
            except Exception:
                pass  # senter may already be active in some model versions
            # Verify that sentence segmentation is actually available.
            # If neither senter nor parser is present, doc.sents yields the whole
            # text as one sentence — silently producing one massive chunk per PDF.
            if not _nlp.has_pipe("senter") and not _nlp.has_pipe("parser"):
                _spacy_unavailable = True
                _log.warning(
                    "spacy_no_sentence_segmenter_using_regex_fallback "
                    "(model has neither senter nor parser — "
                    "run: python -m spacy download en_core_web_sm)"
                )
        except Exception:
            _spacy_unavailable = True
            _log.warning("spacy_unavailable_using_regex_fallback "
                         "(install: python -m spacy download en_core_web_sm)")
        if _spacy_unavailable:
            parts = re.split(r"(?<=[.!?])\s+", text)
            return [p.strip() for p in parts if p.strip()]
    try:
        doc = _nlp(text)
        return [sent.text.strip() for sent in doc.sents if sent.text.strip()]
    except Exception:
        parts = re.split(r"(?<=[.!?])\s+", text)
        return [p.strip() for p in parts if p.strip()]


_PDF_OCR_CHAR_THRESHOLD = 80   # pages with fewer chars after extraction get OCR'd


def _ocr_page(page) -> str:
    """Try to extract text from a fitz page via OCR. Returns '' if OCR not available."""
    try:
        import pytesseract
        from PIL import Image
        import io
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img, config="--psm 6")
    except Exception:
        return ""


def _load_pdf_full(file_path):
    """Full pipeline: pymupdf + noise + sentence chunk. Fallbacks kept.

    For pages where fitz extracts fewer than _PDF_OCR_CHAR_THRESHOLD characters
    (common for dashboard screenshots or image-only slides), an OCR pass is
    attempted via pytesseract if it is installed. Pages where neither fitz nor
    OCR yields text are silently skipped.
    """
    try:
        import fitz
        pages = []
        raw_pages = []
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text = page.get_text("text")
                if len(text.strip()) < _PDF_OCR_CHAR_THRESHOLD:
                    ocr_text = _ocr_page(page)
                    if ocr_text.strip():
                        text = ocr_text
                if text.strip():
                    raw_pages.append(text)
                    pages.append(noise_filter(text))
        full_text = "\n\n".join(pages)
        chunks = sentence_chunk(full_text, size=256, overlap=64)
        if not chunks and raw_pages:
            # noise_filter stripped everything (common for architecture/diagram PDFs
            # with many short labels); retry without filtering
            chunks = sentence_chunk("\n\n".join(raw_pages), size=256, overlap=64)
        if chunks:
            for c in chunks:
                c.metadata["source"] = file_path
                c.metadata.setdefault("doc_type", "qbr_deck")
            return chunks
        # fitz extracted no text at all (image-based PDF) — fall through to PyPDFLoader
    except ImportError:
        pass
    _splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=200)
    docs = _splitter.split_documents(PyPDFLoader(file_path).load())
    for d in docs:
        d.metadata.setdefault("doc_type", "qbr_deck")
    return docs


def _load_docx_full(file_path):
    try:
        import docx
        from docx.oxml.ns import qn as _qn
        from docx.text.paragraph import Paragraph as _DocxParagraph
        from docx.table import Table as _DocxTable

        d = docx.Document(file_path)
        chunks_raw = []
        current_heading = ""
        current_paras: list = []

        def flush():
            if current_paras:
                text = (current_heading + "\n" + "\n".join(current_paras)).strip()
                if text:
                    chunks_raw.append(Document(
                        page_content=text,
                        metadata={"source": file_path, "heading": current_heading},
                    ))

        def _table_to_rows(tbl: "_DocxTable") -> list:
            """Return non-empty rows as pipe-separated strings."""
            rows = []
            for row in tbl.rows:
                cells = [cell.text.strip() for cell in row.cells]
                # Deduplicate merged cells (python-docx repeats merged cell text)
                seen = set()
                unique = []
                for c in cells:
                    if c and c not in seen:
                        seen.add(c)
                        unique.append(c)
                if unique:
                    rows.append(" | ".join(unique))
            return rows

        # Walk body children in document order to maintain heading context for tables
        for child in d.element.body:
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if local == "p":
                para = _DocxParagraph(child, d)
                if para.style.name.startswith("Heading"):
                    flush()
                    current_heading = para.text.strip()
                    current_paras = []
                elif para.text.strip():
                    current_paras.append(para.text.strip())
            elif local == "tbl":
                table = _DocxTable(child, d)
                rows = _table_to_rows(table)
                if rows:
                    # Emit table text as a labelled block under the current heading
                    label = f"{current_heading} [table]" if current_heading else "[table]"
                    current_paras.append(f"{label}\n" + "\n".join(rows))

        flush()
        if chunks_raw:
            docs = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200).split_documents(chunks_raw)
            for doc in docs:
                doc.metadata.setdefault("doc_type", "qbr_deck")
            return docs
    except Exception:
        pass
    docs = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200).split_documents(
        Docx2txtLoader(file_path).load())
    for doc in docs:
        doc.metadata.setdefault("doc_type", "qbr_deck")
    return docs


def _load_html_full(file_path):
    try:
        from langchain_text_splitters import HTMLHeaderTextSplitter
        splitter = HTMLHeaderTextSplitter(headers_to_split_on=[
            ("h1", "Header1"), ("h2", "Header2"), ("h3", "Header3")])
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            html = f.read()
        docs = splitter.split_text(html)
        for d in docs:
            d.metadata.setdefault("source", file_path)
        result = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200).split_documents(docs)
        for d in result:
            d.metadata.setdefault("doc_type", "solution_architecture")
        return result
    except Exception:
        # BeautifulSoup fallback — avoids fetching external resources referenced
        # in the HTML (scripts, stylesheets), which would be a potential SSRF vector.
        from bs4 import BeautifulSoup
        with open(file_path, "r", encoding="utf-8", errors="replace") as _f:
            _html = _f.read()
        _text = BeautifulSoup(_html, "html.parser").get_text(separator="\n")
        _docs = [Document(page_content=_text, metadata={"source": file_path})]
        result = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200).split_documents(_docs)
        for d in result:
            d.metadata.setdefault("doc_type", "solution_architecture")
        return result


def _load_transcript(file_path):
    """Transcript loader with fallback to plain-text when no real speaker labels found.

    The previous version always assumed `.txt` was a speaker-formatted transcript.
    A plain memo `.txt` with no `Speaker: ...` lines parsed as a single turn with
    speaker="Speaker" — semantically meaningless. We now detect that case and fall
    back to PDF-style sentence chunking so plain text still indexes usefully.
    """
    from ingestion.transcript_parser import parse as parse_t
    from ingestion.transcript_chunker import chunk as chunk_t
    turns = parse_t(file_path)
    # Fall back to plain-text only when the entire file parsed as a single turn
    # (no speaker labels found at all). Multi-turn single-speaker transcripts
    # (webinars, lectures) still go through chunk_t — _load_plain_text's
    # noise_filter strips short dialogue lines (<4 words) which destroys content.
    if len(turns) <= 1:
        _log.info("transcript_parser_no_speaker_labels_falling_back_to_plain_text",
                  extra={"file": os.path.basename(file_path),
                         "turn_count": len(turns)})
        return _load_plain_text(file_path)
    return chunk_t(turns, source=file_path)


def _load_ticket(file_path):
    """Ticket loader with fallback to generic-JSON when not ticket-shaped."""
    from ingestion.ticket_parser import parse as parse_t
    from ingestion.ticket_chunker import chunk as chunk_t
    try:
        ticket = parse_t(file_path)
        # Real ticket signal: has subject OR at least one body field.
        # Subject-only tickets (common in Zendesk exports) are valid — rejecting them
        # caused fallback to _load_generic_json, losing all structured metadata.
        if not (ticket.subject or ticket.description or ticket.comments or ticket.resolution):
            raise ValueError("not ticket-shaped (no subject/description/comments/resolution)")
        return chunk_t(ticket, source=file_path)
    except Exception as e:
        _log.info("ticket_parser_unrecognized_falling_back_to_genericjson",
                  extra={"file": os.path.basename(file_path), "error": str(e)})
        return _load_generic_json(file_path)


def _load_plain_text(file_path):
    """Load a .txt file as plain prose — reuse the PDF noise filter + sentence chunker."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    text = noise_filter(text)
    chunks = sentence_chunk(text, size=256, overlap=64)
    for c in chunks:
        c.metadata["source"] = file_path
        c.metadata["doc_type"] = "plain_text"
    return chunks


# ── Document date resolution ─────────────────────────────────────────────────
# A chunk's `doc_date` represents WHEN THE CONTENT IS FROM, distinct from
# WHEN IT WAS UPLOADED. Without this, a transcript uploaded today that describes
# a call from three weeks ago answers "the most recent call" incorrectly.

_FILENAME_DATE_RE = re.compile(r"^(\d{4}-\d{2}-\d{2})")
_CONTENT_DATE_RE  = re.compile(r"\b(\d{4}-\d{2}-\d{2})\b")


def resolve_doc_date(filename: str, contents: bytes, doc_type: str) -> Optional[str]:
    """Return a YYYY-MM-DD content date, or None if no date is recoverable.

    Priority (first hit wins):
      1. Format-specific extraction (content is authoritative for structured types):
         - transcript / plain_text / account_notes: regex over first 1 KB
         - ticket / genericjson: ``created_at`` / ``updated_at`` fields
         - commitment_tracker: latest of all target_date / completed_date values
      2. Filename prefix: e.g. ``2025-03-28_meridian_call.txt``
         (fallback only — user-supplied filenames are frequently inaccurate)
      3. None — caller falls back to upload timestamp

    Used at ingest time to stamp ``doc_date`` on every chunk's metadata so
    retrieval and the LLM can reason about recency.
    """
    # 1a. Text-based content extraction (transcript, account notes)
    if doc_type in ("transcript", "plain_text", "account_notes"):
        try:
            head = contents[:2048].decode("utf-8", errors="ignore")
            m = _CONTENT_DATE_RE.search(head)
            if m:
                return m.group(1)
        except Exception:
            pass

    # 1b. Ticket / generic JSON — use the most recent timestamp field
    if doc_type in ("ticket", "generic_json"):
        try:
            data = json.loads(contents)
            if isinstance(data, dict):
                # Prefer updated_at over created_at — it's more recent
                for field in ("updated_at", "created_at", "created", "date"):
                    ts = data.get(field)
                    if isinstance(ts, str) and len(ts) >= 10 and _CONTENT_DATE_RE.match(ts[:10]):
                        return ts[:10]
        except Exception:
            pass

    # 1c. Commitment tracker JSON — use latest date across all records
    if doc_type == "commitment_tracker":
        try:
            data = json.loads(contents)
            items = data if isinstance(data, list) else data.get("commitments", [])
            dates = []
            for item in items:
                for field in ("current_target_date", "promised_date", "last_updated"):
                    val = item.get(field, "")
                    if isinstance(val, str) and _CONTENT_DATE_RE.match(val[:10] if len(val) >= 10 else ""):
                        dates.append(val[:10])
            if dates:
                return max(dates)
        except Exception:
            pass

    # 2. Filename prefix fallback (user-supplied, may be inaccurate)
    m = _FILENAME_DATE_RE.match(filename or "")
    if m:
        return m.group(1)

    return None


_MEETING_TYPE_KEYWORDS = (
    "qbr", "status-call", "status_call", "incident-review",
    "incident_review", "kickoff", "renewal",
)


def resolve_doc_metadata(filename: str, contents: bytes, doc_type: str) -> dict:
    """Extract structured metadata fields beyond doc_date, keyed by doc_type.

    Called at ingest time so every chunk in a workspace gets uniform metadata
    regardless of which loaders set what. Returned dict is merged (not replaced)
    into existing chunk metadata via setdefault — loaders' own fields win.
    """
    meta: dict = {}
    if not doc_type:
        return meta

    if doc_type == "ticket":
        try:
            data = json.loads(contents)
            if isinstance(data, dict):
                for field in ("reporter", "assignee", "updated_at"):
                    val = data.get(field)
                    if val:
                        meta[field] = str(val)
        except Exception:
            pass

    if doc_type in ("transcript", "account_notes"):
        lower = (filename or "").lower()
        for kw in _MEETING_TYPE_KEYWORDS:
            if kw in lower:
                meta["meeting_type"] = kw.replace("_", "-")
                break
        # author from filename: YYYY-MM-DD_<type>_<author>.txt
        parts = os.path.splitext(filename or "")[0].split("_")
        if len(parts) >= 3:
            meta.setdefault("author", parts[-1])

    if doc_type == "account_notes":
        meta["is_internal"] = "true"

    if doc_type == "commitment_tracker":
        meta["is_structured"] = "true"

    return meta


def _wants_recency(query: str) -> bool:
    """Detect when a query has a recency intent that should boost newer chunks."""
    return bool(_RECENCY_KEYWORDS.search(query))


_RECENCY_KEYWORDS = re.compile(
    r"\b(recent|latest|last (?:call|meeting|week|month|quarter|year)|"
    r"yesterday|today|this (?:week|month|quarter|year)|"
    r"current|new(?:est)?|since|most recent|just|lately)\b",
    re.I,
)


def _recency_boost(docs: List[Document], semantic_weight: float = 0.4) -> List[Document]:
    """Re-rank docs blending semantic rank with recency.

    Newer doc_dates rank higher. ``semantic_weight`` of 0.4 means a 60/40 split
    favoring recency — chosen so an on-topic-but-old chunk doesn't beat a slightly
    less topical but very recent chunk when the user asked for "recent".

    Documents without a parseable doc_date sink to the bottom (epoch fallback).
    """
    if not docs:
        return docs
    parsed = []
    for d in docs:
        ds = d.metadata.get("doc_date") or "1970-01-01"
        try:
            dt = datetime.strptime(ds[:10], "%Y-%m-%d")
        except Exception:
            dt = datetime(1970, 1, 1)
        parsed.append(dt)

    oldest = min(parsed)
    newest = max(parsed)
    span_days = max((newest - oldest).days, 1)
    n = len(docs)
    blended = []
    for i, (doc, dt) in enumerate(zip(docs, parsed)):
        sem_score = (n - i) / n  # rank position → linear score
        recency_score = (dt - oldest).days / span_days
        score = semantic_weight * sem_score + (1.0 - semantic_weight) * recency_score
        blended.append((score, i, doc))  # `i` as stable tiebreak
    blended.sort(key=lambda x: (-x[0], x[1]))
    return [d for _, _, d in blended]


def _walk_json_strings(obj, out=None):
    """Recursively collect meaningful string values from a nested JSON structure.

    Skips strings shorter than 10 chars — single-word enum values ("true", "1",
    "open", "NA") add no retrievable signal and pollute chunk text.
    """
    if out is None:
        out = []
    if isinstance(obj, str):
        if len(obj.strip()) >= 10:
            out.append(obj.strip())
    elif isinstance(obj, dict):
        for v in obj.values():
            _walk_json_strings(v, out)
    elif isinstance(obj, list):
        for v in obj:
            _walk_json_strings(v, out)
    # ints/floats/bools/null: skip
    return out


def _load_generic_json(file_path):
    """Load JSON that isn't ticket-shaped — concatenate string values, sentence-chunk."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        try:
            data = json.load(f)
        except json.JSONDecodeError:
            # Fall back to treating it as plain text rather than raising
            return _load_plain_text(file_path)
    flat = _walk_json_strings(data)
    text = "\n".join(flat)
    chunks = sentence_chunk(text, size=256, overlap=64) if text.strip() else []
    for c in chunks:
        c.metadata["source"] = file_path
        c.metadata["doc_type"] = "generic_json"
    return chunks


_CSV_TERMINAL_STATUSES = frozenset({
    "closed", "resolved", "done", "complete", "completed",
    "fixed", "won't fix", "wontfix", "duplicate", "cancelled", "canceled",
    "rejected", "invalid",
})


def _load_csv_tickets(file_path):
    """Load a CSV export of tickets (Jira/Zendesk) using ticket_csv_parser.

    Each ticket is split by section (header, description, resolution) and long
    sections are further split so the embedding model never receives truncated input.
    """
    from ingestion.ticket_csv_parser import parse_csv
    from ingestion.ticket_chunker import _split_long_text
    try:
        records = parse_csv(file_path)
    except Exception as e:
        _log.warning("ticket_csv_parse_failed fallback_to_plain file=%s error=%s",
                     os.path.basename(file_path), str(e))
        return _load_plain_text(file_path)
    docs = []
    for rec in records:
        base_meta = {
            "source": file_path,
            "doc_type": "ticket",
            "ticket_id": rec.ticket_id or "",
            "status": rec.status or "",
            "priority": rec.priority or "",
            "created_date": rec.created_date or "",
            "doc_date": rec.created_date or "",
            "is_open": str((rec.status or "").lower() not in _CSV_TERMINAL_STATUSES).lower(),
        }
        header = (
            f"Ticket: {rec.ticket_id}\nTitle: {rec.summary}\n"
            f"Status: {rec.status}\nPriority: {rec.priority}"
        )
        # Description — split if long
        if rec.description:
            for seg_i, seg in enumerate(_split_long_text(rec.description)):
                meta = {**base_meta, "section": "description"}
                if seg_i > 0:
                    meta["section_part"] = seg_i
                docs.append(Document(
                    page_content=f"{header}\nDescription: {seg}",
                    metadata=meta,
                ))
        else:
            docs.append(Document(page_content=header, metadata={**base_meta, "section": "description"}))
        # Resolution — split if long
        if rec.resolution:
            for seg_i, seg in enumerate(_split_long_text(rec.resolution)):
                meta = {**base_meta, "section": "resolution"}
                if seg_i > 0:
                    meta["section_part"] = seg_i
                docs.append(Document(
                    page_content=f"Ticket: {rec.ticket_id} Title: {rec.summary}\nResolution: {seg}",
                    metadata=meta,
                ))
    if not docs:
        return _load_plain_text(file_path)
    return docs


def _load_csv_commitments(file_path):
    """Load a CSV export of commitments (Google Sheets) using commitment_parser."""
    from ingestion.commitment_parser import parse_csv_sheets
    try:
        records = parse_csv_sheets(file_path)
    except Exception as e:
        _log.warning("commitment_csv_parse_failed fallback_to_plain file=%s error=%s",
                     os.path.basename(file_path), str(e))
        return _load_plain_text(file_path)

    today_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    docs = []
    for rec in records:
        target = rec.current_target_date or rec.promised_date or ""
        is_overdue = bool(target and getattr(rec, "is_open", True) and target < today_str)
        days_overdue = 0
        if is_overdue:
            try:
                days_overdue = (
                    datetime.strptime(today_str, "%Y-%m-%d") - datetime.strptime(target, "%Y-%m-%d")
                ).days
            except Exception:
                pass

        overdue_line = f"\nOVERDUE by {days_overdue} days (target was {target})" if is_overdue else ""
        text = f"Commitment: {rec.description}{overdue_line}\nStatus: {rec.status}\n"
        if rec.promised_date:
            text += f"Promised: {rec.promised_date}\n"
        if rec.current_target_date:
            text += f"Target: {rec.current_target_date}\n"
        if rec.owner:
            text += f"Owner: {rec.owner}\n"
        doc = Document(
            page_content=text.strip(),
            metadata={
                "source": file_path,
                "doc_type": "commitment_tracker",
                "commitment_id": rec.commitment_id or "",
                "status": rec.status or "open",
                "promised_date": rec.promised_date or "",
                "current_target_date": rec.current_target_date or "",
                "owner": rec.owner or "",
                "is_slipped": str(getattr(rec, "is_slipped", False)).lower(),
                "is_overdue": str(is_overdue).lower(),
                "days_overdue": days_overdue,
                "customer_aware": str(getattr(rec, "customer_aware", False)).lower(),
                # doc_date: current_target_date reflects the commitment's most recent
                # relevant date; falls back to promised_date if no target is set.
                "doc_date": rec.current_target_date or rec.promised_date or "",
            }
        )
        docs.append(doc)
    if not docs:
        return _load_plain_text(file_path)
    return docs


# ── Format-aware ingest strategy registry ─────────────────────────────────────
#
# Single source of truth for "given file extension X, how do we load and index?"
#
# Two key decisions per format:
#   1. WHICH LOADER fires — picked to preserve the format's natural structure.
#      PDF: pymupdf + spaCy sentence boundaries (free-form prose).
#      DOCX: heading-grouped paragraphs (explicit hierarchy).
#      HTML: h1/h2/h3 splits (semantic structure).
#      TXT: speaker-aware turn grouping for transcripts; falls back to plain
#           prose if no speaker labels found (so a generic memo still indexes).
#      JSON: ticket-section parsing for support-ticket shape; falls back to
#            generic-JSON walk for non-ticket shapes.
#
#   2. PARENT-CHILD RE-SPLIT — applied to free-form prose (PDF/DOCX/HTML)
#      where a smaller embedded chunk benefits from a larger surrounding parent
#      window for LLM context. NOT applied to transcripts or tickets, where
#      the format-aware loader already produces "right-sized" chunks (a
#      speaker-turn group, a ticket section). Re-splitting those at character
#      boundaries would bisect a speaker turn or ticket section — losing the
#      attribution that makes them useful to retrieval.
#
# Adding a new format = add one entry here. Don't scatter `if ext == ".x":`
# checks across loader, indexer, and dispatch sites.


@dataclass(frozen=True)
class IngestStrategy:
    """How a document of one format gets loaded, chunked, and indexed."""
    name: str                                          # short label for logs / metadata
    loader: Callable[[str], List[Document]]            # produces format-aware chunks
    use_parent_child_resplit: bool                     # see notes above
    natural_unit: str                                  # what one chunk represents
    rationale: str                                     # why these choices fit this format


_FULL_STRATEGY: Dict[str, IngestStrategy] = {
    ".pdf": IngestStrategy(
        name="pdf",
        loader=_load_pdf_full,
        use_parent_child_resplit=True,
        natural_unit="page-spanning paragraph (sentence-bounded, ~256 words)",
        rationale=(
            "PDF prose varies wildly in section size. pymupdf gives clean text "
            "extraction; noise_filter strips page numbers and headers/footers; "
            "spaCy sentence boundaries prevent mid-sentence cuts. Parent-child "
            "re-split (1600/500) gives precise embedded children + larger parent "
            "context windows for the LLM."
        ),
    ),
    ".docx": IngestStrategy(
        name="docx",
        loader=_load_docx_full,
        use_parent_child_resplit=True,
        natural_unit="heading-grouped paragraph block",
        rationale=(
            "DOCX has explicit Heading-style metadata. Group body paragraphs "
            "under their preceding Heading so the section context is preserved "
            "in chunk metadata. Parent-child re-split for the same reason as PDF."
        ),
    ),
    ".html": IngestStrategy(
        name="html",
        loader=_load_html_full,
        use_parent_child_resplit=True,
        natural_unit="h1/h2/h3 section",
        rationale=(
            "HTML carries semantic hierarchy via h1-h3 tags. HTMLHeaderTextSplitter "
            "preserves it as metadata so retrieval can favor specific sections. "
            "Parent-child for size normalization across short and long sections."
        ),
    ),
    ".txt": IngestStrategy(
        name="transcript",
        # _load_transcript falls back to plain_text when no speaker labels are
        # found — so a memo .txt still indexes (just without speaker metadata).
        loader=_load_transcript,
        use_parent_child_resplit=False,
        natural_unit="per-speaker-turn segment (same-speaker runs merged ≤120 words, prior-turn context prepended)",
        rationale=(
            "DECISION ANCHOR: no character-boundary bisection of speaker turns. "
            "Bisecting at 500 chars produces fragments like 'Sarah: We agreed we'd "
            "ship' — incoherent in citations and mis-attributing tail text to the "
            "wrong speaker. Per-speaker-segment chunks fix the speaker-semantic "
            "dilution problem: one chunk = one speaker's utterance, so the embedding "
            "captures a single person's meaning rather than a blend of all speakers "
            "in a 200-word window. Short consecutive same-speaker turns are merged "
            "(≤120 words) to avoid noisy one-line chunks. A leading [Prior: ...] "
            "context line in each chunk gives the LLM enough surrounding context "
            "for accurate attribution. use_parent_child_resplit=False: speaker-segment "
            "chunks are already the right size and are their own natural context unit."
        ),
    ),
    ".json": IngestStrategy(
        name="ticket",
        # _load_ticket falls back to genericjson for non-ticket-shaped JSON.
        loader=_load_ticket,
        use_parent_child_resplit=False,
        natural_unit="ticket section (description / single comment / resolution)",
        rationale=(
            "DECISION ANCHOR: same citation-quality reasoning as transcripts. "
            "Each ticket section is a coherent unit with its own author and "
            "timestamp. A query like 'what was the resolution?' should land "
            "directly on the resolution section, not a 500-char fragment of "
            "it. Re-splitting fragments the section identity that makes ticket "
            "retrieval useful — embedding model speed has no bearing on this."
        ),
    ),
    ".csv": IngestStrategy(
        name="csv",
        loader=_load_csv_tickets,
        use_parent_child_resplit=False,
        natural_unit="one ticket or commitment row",
        rationale=(
            "CSV exports (Jira/Zendesk tickets, Google Sheets commitments) are "
            "already structured. Each row is a self-contained unit. Splitting at "
            "character boundaries would break the row structure. doc_type override "
            "routes commitment CSVs to _load_csv_commitments via the caller."
        ),
    ),
}

# Future hook: a "baseline" strategy for the chunking-ablation experiments lives
# in experiment_kit/legacy_loaders.py — it imports the same _load_*_baseline
# functions but doesn't go through this registry at runtime.


def lookup_strategy(file_path: str) -> Optional[IngestStrategy]:
    """Return the IngestStrategy for a file, or None if the format is unsupported."""
    ext = os.path.splitext(file_path)[1].lower()
    return _FULL_STRATEGY.get(ext)


def load_and_split_document(file_path):
    """Format-aware loader dispatch.

    Selects an :class:`IngestStrategy` from the registry by file extension and
    runs its loader. Logs the strategy name and chunk count so operators can
    verify the right path fired without instrumentation.

    The returned chunks are the format-aware loader's output BEFORE any
    parent-child re-split. ``index_document_to_chroma`` reads
    ``strategy.use_parent_child_resplit`` to decide whether to fan out into
    smaller embedded children + larger parent context windows.

    Raises ValueError on an unrecognized extension.
    """
    fname = os.path.basename(file_path)
    strategy = lookup_strategy(file_path)
    if strategy is None:
        raise ValueError(f"Unsupported file extension: {file_path}")

    docs = strategy.loader(file_path)
    sample_doc_type = docs[0].metadata.get("doc_type", "") if docs else ""
    _log.info(
        "load_and_split_document strategy=%s chunks=%d ext=%s file=%s doc_type=%s "
        "natural_unit=%s parent_child_resplit=%s",
        strategy.name, len(docs), os.path.splitext(file_path)[1].lower(),
        fname, sample_doc_type, strategy.natural_unit,
        strategy.use_parent_child_resplit,
    )
    return docs


# ── Indexing: parent-child for "full", flat for "baseline" and "sentence" ────
# Parent-child sizing rationale:
#   Child (800 chars ≈ 160 words ≈ 200 tokens): embedded unit — sits in the
#     optimal dense-retrieval range for text-embedding-004 (100–250 tokens).
#   Parent (2400 chars ≈ 480 words): LLM context unit — 3:1 ratio gives the
#     model a full three-child window of surrounding text for synthesis questions.
#     The previous 1600-char parent (2:1 ratio) added only one extra child-width
#     of context, which was insufficient for multi-paragraph answers.
#   Child overlap (150 chars ≈ 18.75%): sentences that span a chunk boundary
#     appear in both adjacent children, preserving retrieval signal at the seam.
_PARENT_SPLITTER = RecursiveCharacterTextSplitter(chunk_size=2400, chunk_overlap=200)
_CHILD_SPLITTER  = RecursiveCharacterTextSplitter(chunk_size=800,  chunk_overlap=150)


_CONTEXTUAL_RETRIEVAL = os.getenv("CONTEXTUAL_RETRIEVAL", "").lower() in ("1", "true", "yes")

_CONTEXT_PROMPT = """<document>
{document}
</document>

Here is the chunk we want to situate within the whole document:
<chunk>
{chunk}
</chunk>

Give a short (1-2 sentence) context to situate this chunk within the overall
document, for the purpose of improving search retrieval. Answer only with the
context itself — no preamble, no prefix, just the description."""


def _contextualize_chunks(chunks, filename):
    """Prepend an LLM-generated context sentence to each chunk.

    Implements Anthropic's contextual retrieval (Sep 2024): each chunk is
    situated within the whole document before embedding, so that cross-chunk
    references ("the fix" / "this customer" / "as mentioned above") and
    section context are recovered. Reported +35-49% retrieval accuracy on
    Anthropic's benchmarks. Enabled via CONTEXTUAL_RETRIEVAL=1.

    Falls back to the original chunk silently on any per-chunk LLM error so
    ingestion is never blocked by a transient model failure.
    """
    if not chunks:
        return chunks
    from langchain_google_genai import ChatGoogleGenerativeAI
    from langchain_core.documents import Document
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        google_api_key=os.getenv("GOOGLE_API_KEY"),
        temperature=0,
    )
    # Reconstruct full document once before the loop. Capped so we stay well
    # below Gemini's input-token budget even for very long documents.
    # Doing this inside the loop would rebuild the same string for every chunk.
    full_doc = "\n\n".join(c.page_content for c in chunks)[:30000]
    enriched = []
    for chunk in chunks:
        try:
            prompt = _CONTEXT_PROMPT.format(document=full_doc, chunk=chunk.page_content)
            context = llm.invoke(prompt).content.strip()
            if not context:
                # Empty LLM response — treat as a failed contextualization so
                # downstream strippers don't get confused by a missing prefix.
                enriched.append(chunk)
                continue
            enriched.append(Document(
                page_content=f"[Context: {context}]\n\n{chunk.page_content}",
                metadata={
                    **chunk.metadata,
                    "has_context_prefix": True,
                    "context_text": context,      # kept for debugging/transparency
                },
            ))
        except Exception as e:
            _log.warning("contextualize_chunk_skipped",
                         extra={"doc_filename": filename, "error": str(e)})
            enriched.append(chunk)
    return enriched


def _existing_workspace_contextual_flag(user_id: str):
    """Read back the contextual-retrieval flag from an existing chunk in this
    workspace, or return None if the workspace has no chunks yet.

    Used to detect mixed-embedding scenarios: if a user ingests some docs with
    CONTEXTUAL_RETRIEVAL=0 and later uploads more with =1, the vector space is
    silently inconsistent (some chunks have an LLM-generated prefix baked into
    their embedding; others don't). Retrieval quality degrades without warning.
    """
    try:
        existing = vectorstore._collection.get(
            where={"user_id": user_id}, limit=1, include=["metadatas"]
        )
        metas = (existing or {}).get("metadatas") or []
        if not metas:
            return None
        return bool(metas[0].get("ingest_contextual_retrieval", False))
    except Exception:
        return None


def _warn_if_mixed_contextual(user_id: str) -> None:
    existing_flag = _existing_workspace_contextual_flag(user_id)
    current_flag = bool(_CONTEXTUAL_RETRIEVAL)
    if existing_flag is not None and existing_flag != current_flag:
        _log.warning(
            "mixed_contextual_workspace user_id=%s existing_flag=%s current_flag=%s "
            "(retrieval quality will be inconsistent — wipe workspace and re-ingest "
            "all docs under a single flag to recover)",
            user_id, existing_flag, current_flag,
        )


def index_document_to_chroma(file_path, file_id, user_id="default", filename=None, doc_type=None):
    """Index a document into Chroma. Returns a summary dict on success, None on failure.

    The summary describes which path was taken so callers (and operators reading
    logs) can verify the right loader and indexing strategy fired. Shape::

        {
          "loader_path": "pdf_full" | "transcript" | "ticket" | ... ,
          "raw_chunks": int,           # chunks from the format-aware loader
          "parent_chunks": int,        # parents stored (== raw_chunks for txt/json)
          "child_chunks": int,         # children embedded (== raw_chunks for txt/json)
          "parent_child_split": bool,  # whether parent-child re-split was applied
          "doc_date": "YYYY-MM-DD",
          "doc_date_source": "content" | "upload_fallback",
          "contextual_retrieval": bool,
        }
    """
    fname = filename or os.path.basename(file_path)
    _warn_if_mixed_contextual(user_id)
    try:
        # commitment_tracker / transcript: bypass the extension-based dispatch.
        # .json would otherwise route to the ticket loader for both doc types.
        if doc_type == "commitment_tracker":
            ext = os.path.splitext(file_path)[1].lower()
            if ext == ".csv":
                raw_docs = _load_csv_commitments(file_path)
            else:
                from ingestion.commitment_parser import parse as _parse_commit
                from ingestion.commitment_chunker import chunk as _chunk_commit
                _commits = _parse_commit(file_path)
                raw_docs = _chunk_commit(
                    _commits, source=fname,
                    today=datetime.now(timezone.utc).strftime("%Y-%m-%d"),
                )
        elif doc_type == "transcript":
            # Always use the transcript loader regardless of extension so that
            # Otter.ai JSON exports (.json) are parsed as transcripts, not tickets.
            raw_docs = _load_transcript(file_path)
        else:
            raw_docs = load_and_split_document(file_path)
    except Exception as e:
        _log.error("indexing_load_failed",
                   extra={"doc_filename": fname, "file_id": file_id, "error": str(e)})
        return None
    if not raw_docs:
        _log.warning("indexing_skipped_no_chunks_extracted",
                     extra={"doc_filename": fname, "file_id": file_id})
        return None

    # Stamp explicit doc_type from upload request onto every chunk.
    # This overrides the format-level default set by individual loaders
    # (e.g. "qbr_deck" for PDF → becomes "solution_architecture" when caller says so).
    if doc_type:
        for d in raw_docs:
            d.metadata["doc_type"] = doc_type

    # Stamp is_latest_version=1 on all new chunks.
    # db_utils.set_latest_version_flag() will flip old chunks to 0 for state docs;
    # snapshot docs (transcripts) keep all versions at 1 (each is historical record).
    for d in raw_docs:
        d.metadata["is_latest_version"] = 1

    # Resolve content date (transcript date, ticket created_at, filename prefix).
    # Falls back to today (upload date) if no content date can be extracted.
    try:
        with open(file_path, "rb") as _fh:
            _contents = _fh.read()
    except Exception:
        _contents = b""
    effective_doc_type = doc_type or (raw_docs[0].metadata.get("doc_type", "") if raw_docs else "")
    resolved_date = resolve_doc_date(fname, _contents, effective_doc_type)

    # Stamp additional structured metadata (reporter, assignee, meeting_type, etc.)
    extra_meta = resolve_doc_metadata(fname, _contents, effective_doc_type)
    if extra_meta:
        for d in raw_docs:
            for k, v in extra_meta.items():
                d.metadata.setdefault(k, v)
    fallback_date = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    effective_date = resolved_date or fallback_date
    date_source = "content" if resolved_date else "upload_fallback"

    def _stamp_dates(docs):
        for d in docs:
            d.metadata.setdefault("doc_date", effective_date)
            d.metadata.setdefault("doc_date_source", date_source)

    try:
        if CHUNKING_MODE == "full":
            # Parent-child pipeline. The decision of WHETHER to apply the
            # parent-child re-split lives in the FORMAT_STRATEGY registry
            # (see lookup_strategy). PDF/DOCX/HTML re-split (free-form prose
            # benefits from precise children + larger parent context windows);
            # transcripts and tickets do NOT (their format-aware chunks are
            # already the natural unit and re-splitting would bisect speaker
            # turns / ticket sections).
            strategy = lookup_strategy(file_path)
            use_parent_child_split = (
                strategy.use_parent_child_resplit if strategy else True
            )

            if use_parent_child_split:
                parent_docs = _PARENT_SPLITTER.split_documents(raw_docs)
                for i, p in enumerate(parent_docs):
                    p.metadata["parent_chunk_id"] = f"p_{file_id}_{i}"
                    p.metadata["file_id"] = file_id
                    p.metadata["user_id"] = user_id
                    p.metadata["filename"] = fname
                _stamp_dates(parent_docs)
                parent_ids = [p.metadata["parent_chunk_id"] for p in parent_docs]
                parent_store.add_documents(parent_docs, ids=parent_ids)

                child_docs = []
                for p in parent_docs:
                    for c in _CHILD_SPLITTER.split_documents([p]):
                        c.metadata["file_id"] = file_id
                        c.metadata["user_id"] = user_id
                        c.metadata["parent_chunk_id"] = p.metadata["parent_chunk_id"]
                        c.metadata["filename"] = fname
                        child_docs.append(c)
                _stamp_dates(child_docs)
            else:
                # Transcript / ticket / plain-text / generic-json: format-aware
                # chunks are the natural unit AND the LLM-facing context unit.
                # There's no larger parent to swap them for, so:
                #   - Children go into the vector store as usual (for retrieval).
                #   - We DO NOT add anything to parent_store. Skipping this saves
                #     50% of the embedding work for these formats — embedding a
                #     copy of the same content into a second collection costs CPU
                #     time and produces an identical vector that retrieval would
                #     never differentiate from the child's.
                #   - We DO NOT set `parent_chunk_id` on children. retrieve_node's
                #     fetch_parents() falls back to returning the children
                #     unchanged when no parent_chunk_id is present — which is
                #     exactly what we want here (the children ARE the context).
                child_docs = []
                for c in raw_docs:
                    c.metadata["file_id"] = file_id
                    c.metadata["user_id"] = user_id
                    c.metadata["filename"] = fname
                    child_docs.append(c)
                _stamp_dates(child_docs)
                # parent_docs intentionally left empty for the summary; no
                # parent_store writes for txt/json.
                parent_docs = []

            # Contextual retrieval (Anthropic, Sep 2024): prepend an LLM-
            # generated doc-level context to each child before embedding.
            if _CONTEXTUAL_RETRIEVAL:
                _log.info("contextualize_started",
                          extra={"doc_filename": fname, "child_count": len(child_docs)})
                child_docs = _contextualize_chunks(child_docs, fname)
            for c in child_docs:
                c.metadata["ingest_contextual_retrieval"] = bool(_CONTEXTUAL_RETRIEVAL)
            vectorstore.add_documents(child_docs)
            summary = {
                "loader_path": raw_docs[0].metadata.get("doc_type") or "unknown",
                "raw_chunks": len(raw_docs),
                "parent_chunks": len(parent_docs),
                "child_chunks": len(child_docs),
                "parent_child_split": use_parent_child_split,
                "doc_date": effective_date,
                "doc_date_source": date_source,
                "contextual_retrieval": bool(_CONTEXTUAL_RETRIEVAL),
            }
        else:
            # Flat: raw_docs go straight into the child store. No parent lookup.
            for d in raw_docs:
                d.metadata["file_id"] = file_id
                d.metadata["user_id"] = user_id
                d.metadata["filename"] = fname
                # No parent_chunk_id → fetch_parents becomes a no-op passthrough
            _stamp_dates(raw_docs)
            if _CONTEXTUAL_RETRIEVAL:
                _log.info("contextualize_started",
                          extra={"doc_filename": fname, "child_count": len(raw_docs)})
                raw_docs = _contextualize_chunks(raw_docs, fname)
            for d in raw_docs:
                d.metadata["ingest_contextual_retrieval"] = bool(_CONTEXTUAL_RETRIEVAL)
            vectorstore.add_documents(raw_docs)
            summary = {
                "loader_path": raw_docs[0].metadata.get("doc_type") or "unknown",
                "raw_chunks": len(raw_docs),
                "parent_chunks": 0,
                "child_chunks": len(raw_docs),
                "parent_child_split": False,
                "doc_date": effective_date,
                "doc_date_source": date_source,
                "contextual_retrieval": bool(_CONTEXTUAL_RETRIEVAL),
            }
        _log.info(
            "indexing_complete file_id=%s filename=%s loader=%s raw_chunks=%d "
            "parent_chunks=%d child_chunks=%d parent_child_split=%s "
            "doc_date=%s (source=%s) contextual=%s",
            file_id, fname,
            summary["loader_path"], summary["raw_chunks"],
            summary["parent_chunks"], summary["child_chunks"],
            summary["parent_child_split"],
            summary["doc_date"], summary["doc_date_source"],
            summary["contextual_retrieval"],
        )
        return summary
    except Exception as e:
        # Surface the actual exception text in the message body, not just
        # extra={...}. Python's default formatter only emits message + level,
        # so debug info hidden in `extra` is invisible without a JSON formatter.
        # Common cause to watch for here: dimension mismatch when the embedding
        # model was changed without wiping the existing Chroma collection.
        _log.exception(
            "indexing_failed file=%s file_id=%s error=%s",
            fname, file_id, str(e),
        )
        return None


def delete_doc_from_chroma(file_id, user_id="default"):
    try:
        cond = {"$and": [{"file_id": {"$eq": file_id}}, {"user_id": {"$eq": user_id}}]}
        vectorstore._collection.delete(where=cond)
        parent_store._collection.delete(where=cond)
        return True
    except Exception as e:
        _log.error("delete_doc_from_chroma_failed",
                   extra={"file_id": file_id, "user_id": user_id, "error": str(e)})
        return False


def get_latest_chunks_by_doctype(customer_id: str, doc_type: str) -> List[Document]:
    """Return all chunks from the most recently uploaded file of doc_type for this customer.

    Uses the is_latest_version=1 filter stamped by set_latest_version_flag() in db_utils.
    Returns [] if no chunks found or vectorstore is unhealthy.
    """
    try:
        result = vectorstore._collection.get(
            where={
                "$and": [
                    {"user_id": {"$eq": customer_id}},
                    {"doc_type": {"$eq": doc_type}},
                    {"is_latest_version": {"$eq": 1}},
                ]
            },
            include=["documents", "metadatas"],
        )
        docs_raw = result.get("documents") or []
        metas_raw = result.get("metadatas") or []
        return [
            Document(page_content=docs_raw[i], metadata=metas_raw[i] or {})
            for i in range(len(docs_raw))
        ]
    except Exception as e:
        _log.error("get_latest_chunks_by_doctype_failed customer=%s doc_type=%s error=%s",
                   customer_id, doc_type, str(e))
        return []


def get_chunks_since_date(customer_id: str, since_date: str,
                          exclude_doc_types: tuple = ("ticket",)) -> List[Document]:
    """Return chunks with doc_date >= since_date for this customer.

    Used by recent_changes_node to implement "since last call" window.
    Excludes ticket doc_types by default (open items are handled separately).

    Chroma's $gte operator on string doc_date inside a $and compound filter fails
    silently — returns [] rather than raising. The date filter is applied in Python
    after fetching all is_latest_version chunks for this customer.
    """
    try:
        filters: list = [
            {"user_id": {"$eq": customer_id}},
            {"is_latest_version": {"$eq": 1}},
        ]
        if exclude_doc_types:
            for dt in exclude_doc_types:
                filters.append({"doc_type": {"$ne": dt}})
        where = {"$and": filters}
        result = vectorstore._collection.get(
            where=where,
            include=["documents", "metadatas"],
        )
        docs_raw = result.get("documents") or []
        metas_raw = result.get("metadatas") or []
        all_docs = [
            Document(page_content=docs_raw[i], metadata=metas_raw[i] or {})
            for i in range(len(docs_raw))
        ]
        # YYYY-MM-DD strings sort lexicographically — safe for Python string comparison
        return [
            d for d in all_docs
            if (d.metadata.get("doc_date") or "1970-01-01") >= since_date
        ]
    except Exception as e:
        _log.error("get_chunks_since_date_failed customer=%s since=%s error=%s",
                   customer_id, since_date, str(e))
        return []


def get_recent_resolved_tickets(customer_id: str, since_date: str) -> List[Document]:
    """Return closed/resolved ticket chunks whose updated_at falls on or after since_date.

    Used by recent_changes_node to enumerate tickets that closed in the window
    rather than relying on retrieval to discover them.
    """
    _terminal_list = ["closed", "resolved", "done", "complete", "completed", "fixed"]
    try:
        result = vectorstore._collection.get(
            where={"$and": [
                {"user_id": {"$eq": customer_id}},
                {"doc_type": {"$eq": "ticket"}},
                {"is_latest_version": {"$eq": 1}},
                {"status": {"$in": _terminal_list}},
            ]},
            include=["documents", "metadatas"],
        )
        docs = [
            Document(page_content=result["documents"][i], metadata=result["metadatas"][i] or {})
            for i in range(len(result.get("documents") or []))
        ]
        return [
            d for d in docs
            if (d.metadata.get("updated_at") or d.metadata.get("doc_date") or "") >= since_date
        ]
    except Exception as e:
        _log.error("get_recent_resolved_tickets_failed customer=%s error=%s", customer_id, str(e))
        return []


def get_recent_completed_commitments(customer_id: str, since_date: str) -> List[Document]:
    """Return commitment chunks with a delivered status and doc_date on or after since_date.

    Used by recent_changes_node to enumerate commitments completed in the window.
    """
    _delivered = {"delivered", "done", "complete", "completed", "closed", "resolved", "fixed"}
    try:
        result = vectorstore._collection.get(
            where={"$and": [
                {"user_id": {"$eq": customer_id}},
                {"doc_type": {"$eq": "commitment_tracker"}},
                {"is_latest_version": {"$eq": 1}},
            ]},
            include=["documents", "metadatas"],
        )
        docs = [
            Document(page_content=result["documents"][i], metadata=result["metadatas"][i] or {})
            for i in range(len(result.get("documents") or []))
        ]
        return [
            d for d in docs
            if (d.metadata.get("commitment_status") or d.metadata.get("status") or "").lower()
               in _delivered
            and (d.metadata.get("doc_date") or "") >= since_date
        ]
    except Exception as e:
        _log.error("get_recent_completed_commitments_failed customer=%s error=%s", customer_id, str(e))
        return []


def demote_old_versions_in_chroma(
    customer_id: str, doc_type: str, new_file_id: int, filename: str = ""
) -> int:
    """Set is_latest_version=0 in Chroma for all chunks of this customer+doc_type
    that belong to a prior upload (file_id != new_file_id).

    When filename is provided the demotion is scoped to that filename only.
    This is critical for doc types where multiple independent files coexist
    (e.g. one JSON file per ticket) — without filename scoping, uploading
    TICK-M002 would demote TICK-M001 to is_latest_version=0.

    Called immediately after set_latest_version_flag() so both SQLite and Chroma
    agree on which version is current. Returns the number of chunks demoted.
    """
    demoted = 0
    for store in (vectorstore, parent_store):
        try:
            where_clauses: list = [
                {"user_id": {"$eq": customer_id}},
                {"doc_type": {"$eq": doc_type}},
                {"is_latest_version": {"$eq": 1}},
                {"file_id": {"$ne": new_file_id}},
            ]
            if filename:
                where_clauses.append({"filename": {"$eq": filename}})
            result = store._collection.get(
                where={"$and": where_clauses},
                include=["metadatas"],
            )
            ids = result.get("ids") or []
            metas = result.get("metadatas") or []
            if not ids:
                continue
            updated = [{**m, "is_latest_version": 0} for m in metas]
            store._collection.update(ids=ids, metadatas=updated)
            demoted += len(ids)
        except Exception as e:
            _log.error("demote_old_versions_failed store=%s customer=%s doc_type=%s error=%s",
                       store._collection.name if hasattr(store, "_collection") else "?",
                       customer_id, doc_type, str(e))
    if demoted:
        _log.info("demoted_old_versions customer=%s doc_type=%s count=%d",
                  customer_id, doc_type, demoted)
    return demoted


def get_account_health_metrics(customer_id: str) -> dict:
    """Compute account health KPIs from chunk metadata. Zero LLM calls.

    Aggregates open ticket counts by priority and commitment slip/overdue
    rates from the metadata stamped at ingest time (is_open, is_slipped,
    target_date, priority). Returns raw counts plus a 0-100 health score.
    """
    from datetime import date
    today = date.today().isoformat()

    # ── Tickets ──────────────────────────────────────────────────────────────
    ticket_chunks = get_latest_chunks_by_doctype(customer_id, "ticket")
    open_p0 = 0
    open_p1 = 0
    seen_tickets: set = set()
    for doc in ticket_chunks:
        meta = doc.metadata
        ticket_key = meta.get("ticket_id") or meta.get("file_id") or doc.page_content[:60]
        if ticket_key in seen_tickets:
            continue
        seen_tickets.add(ticket_key)
        is_open = meta.get("is_open") in ("true", True, "1")
        if not is_open:
            continue
        priority = (meta.get("priority") or "").lower().replace("-", "").replace(" ", "")
        if priority in ("p0", "critical", "urgent", "blocker"):
            open_p0 += 1
        elif priority in ("p1", "high"):
            open_p1 += 1

    # ── Commitments ──────────────────────────────────────────────────────────
    commit_chunks = get_latest_chunks_by_doctype(customer_id, "commitment_tracker")
    total_commitments = 0
    open_commitments = 0
    overdue_commitments = 0
    slipped_commitments = 0
    for doc in commit_chunks:
        meta = doc.metadata
        total_commitments += 1
        if meta.get("is_slipped") in ("true", True, "1"):
            slipped_commitments += 1
        is_open = meta.get("is_open") in ("true", True, "1")
        if is_open:
            open_commitments += 1
            if (meta.get("is_overdue") or "").lower() == "true":
                overdue_commitments += 1
            else:
                # fallback for chunks ingested before is_overdue was stamped
                target = (meta.get("current_target_date") or meta.get("target_date") or "").strip()
                if target and target < today:
                    overdue_commitments += 1

    slip_rate = slipped_commitments / total_commitments if total_commitments > 0 else 0.0

    # ── Score ─────────────────────────────────────────────────────────────────
    # Each open P0 is the most severe signal — a single P0 alone should push
    # an account into "At Risk". P1s are significant but not catastrophic.
    # Overdue commitments are weighted heavily; slip rate less so (a commitment
    # can slip without being overdue yet).
    score = 100
    score -= open_p0 * 25
    score -= open_p1 * 8
    score -= overdue_commitments * 12
    score -= round(slip_rate * 20)
    score = max(0, score)

    if score >= 75:
        band = "Healthy"
    elif score >= 45:
        band = "At Risk"
    else:
        band = "Critical"

    return {
        "open_p0_count": open_p0,
        "open_p1_count": open_p1,
        "overdue_commitment_count": overdue_commitments,
        "total_open_commitments": open_commitments,
        "slipped_commitment_count": slipped_commitments,
        "total_commitments": total_commitments,
        "commitment_slip_rate": round(slip_rate, 2),
        "health_score": score,
        "health_band": band,
    }


def get_person_relevant_chunks(
    customer_id: str,
    doc_types: List[str],
    since_date: str = "",
    max_results: int = 80,
) -> List[Document]:
    """Fetch narrative chunks for a customer filtered by doc_type at the Chroma layer.

    Chroma's $gte operator on string doc_date inside a $and compound filter fails
    silently (returns [] instead of raising). The date filter is therefore applied
    in Python after the Chroma fetch — same pattern as get_chunks_since_date().

    The doc_type $in filter IS applied at the Chroma layer, bounding the result set
    to narrative doc types (transcript, account_notes, qbr_deck) before Python sees it.
    Returns docs sorted by doc_date descending, capped at max_results.
    """
    try:
        result = vectorstore._collection.get(
            where={"$and": [
                {"user_id": {"$eq": customer_id}},
                {"doc_type": {"$in": doc_types}},
                {"is_latest_version": {"$eq": 1}},
            ]},
            include=["documents", "metadatas"],
        )
        docs = [
            Document(page_content=result["documents"][i], metadata=result["metadatas"][i] or {})
            for i in range(len(result.get("documents") or []))
        ]
        # Apply date filter in Python — Chroma $gte on string fields fails silently
        if since_date:
            docs = [d for d in docs if (d.metadata.get("doc_date") or "1970-01-01") >= since_date]
        docs.sort(key=lambda d: d.metadata.get("doc_date") or "1970-01-01", reverse=True)
        return docs[:max_results]
    except Exception as e:
        _log.warning(
            "get_person_relevant_chunks_failed customer=%s error=%s",
            customer_id, str(e),
        )
        return []


def structured_metadata_retrieve(
    customer_id: str,
    doc_type: str,
    extra_filters: Optional[List[dict]] = None,
    max_results: int = 30,
) -> List[Document]:
    """Retrieve chunks by metadata filters without semantic search.

    Runs entirely at the ChromaDB layer — no embeddings, no reranking.
    Used by retrieve_node for structured queries about commitments or tickets
    where semantic ranking would miss records that don't score well against
    the query string but match precisely on metadata predicates (status,
    priority, date ranges, is_overdue, etc.).

    extra_filters: list of Chroma where-clause dicts ANDed with the base
    customer + doc_type + is_latest_version filter.
    Returns up to max_results Documents sorted by doc_date descending.
    """
    try:
        where_clauses: list = [
            {"user_id": {"$eq": customer_id}},
            {"doc_type": {"$eq": doc_type}},
            {"is_latest_version": {"$eq": 1}},
        ]
        if extra_filters:
            where_clauses.extend(extra_filters)
        result = vectorstore._collection.get(
            where={"$and": where_clauses},
            include=["documents", "metadatas"],
        )
        docs = [
            Document(page_content=result["documents"][i], metadata=result["metadatas"][i] or {})
            for i in range(len(result.get("documents") or []))
        ]
        docs.sort(key=lambda d: d.metadata.get("doc_date") or "1970-01-01", reverse=True)
        return docs[:max_results]
    except Exception as e:
        _log.warning(
            "structured_metadata_retrieve_failed customer=%s doc_type=%s error=%s",
            customer_id, doc_type, str(e),
        )
        return []
